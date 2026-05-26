"""
Eczanem Bitirme Tezi — Word (.docx) üretici
Fırat Üniversitesi Bilgisayar Mühendisliği yazım kurallarına uygundur.
Çalıştır: python generate_thesis.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ──────────────────────────────────────────────
# Yardımcı fonksiyonlar
# ──────────────────────────────────────────────


def set_page_margins(section):
    """Sol 3.5 cm, Üst/Alt 3 cm, Sağ 2.5 cm"""
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)


def set_run_font(run, size_pt=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    # XML seviyesinde de font adını zorla (Doğu dilleri için)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rPr.insert(0, rFonts)


def add_paragraph(
    doc,
    text="",
    bold=False,
    italic=False,
    size=11,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_before=0,
    space_after=6,
    line_spacing=1.5,
    keep_with_next=False,
):
    """Genel paragraf ekler."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(line_spacing * 11)  # 1.5 × 11pt satır
    pf.keep_with_next = keep_with_next
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=size, bold=bold, italic=italic)
    return p


def add_heading1(doc, text):
    """1. derece başlık: BÜYÜK HARF, KALIN, yeni sayfa."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(22)  # iki satır boşluk ≈ 2×11pt
    pf.line_spacing = Pt(16.5)
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    set_run_font(run, size_pt=11, bold=True)
    return p


def add_heading2(doc, text):
    """2. derece başlık: Her Kelime Büyük, kalın."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(11)
    pf.space_after = Pt(5.5)
    pf.line_spacing = Pt(16.5)
    pf.keep_with_next = True
    # Title-case çeviri (Türkçe büyük harf)
    title_text = " ".join(w.capitalize() for w in text.split())
    run = p.add_run(title_text)
    set_run_font(run, size_pt=11, bold=True)
    return p


def add_bold_inline(paragraph, label, rest, size=11):
    """Aynı paragrafta kalın etiket + normal metin."""
    r1 = paragraph.add_run(label)
    set_run_font(r1, size_pt=size, bold=True)
    r2 = paragraph.add_run(rest)
    set_run_font(r2, size_pt=size)


def add_body(doc, text):
    return add_paragraph(
        doc,
        text,
        size=11,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=0,
        space_after=6,
        line_spacing=1.5,
    )


def add_bullet(doc, text, level=0):
    """Madde işaretli liste satırı."""
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing = Pt(16.5)
    run = p.add_run(text)
    set_run_font(run, size_pt=11)
    return p


def add_table(doc, headers, rows, caption_text=""):
    """Başlıklı tablo ekler."""
    if caption_text:
        cp = add_paragraph(
            doc,
            caption_text,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=6,
            space_after=3,
            line_spacing=1.0,
        )

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Başlık satırı
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size_pt=11, bold=True)
        # Gri arka plan
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    # Veri satırları
    for ri, row_data in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = trow.cells[ci]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(cell_text)
            set_run_font(run, size_pt=10)

    # Tablo sonrası boşluk
    add_paragraph(doc, space_before=6, space_after=0)
    return table


def add_page_number(doc):
    """Her bölüm için sayfa numarası (alt orta)."""
    section = doc.sections[-1]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run()
    set_run_font(run, size_pt=11)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ──────────────────────────────────────────────
# Belge oluşturma
# ──────────────────────────────────────────────

doc = Document()

# Varsayılan stil
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

# Kenar boşlukları
for sec in doc.sections:
    set_page_margins(sec)

add_page_number(doc)

# ══════════════════════════════════════════════
# DIŞ KAPAK
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(6)
for line in ["T.C.", "FIRAT ÜNİVERSİTESİ", "MÜHENDİSLİK FAKÜLTESİ"]:
    run = p.add_run(line + "\n")
    set_run_font(run, size_pt=12, bold=True)

add_paragraph(doc, space_before=36, space_after=36)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run("YAPAY ZEKÂ DESTEKLİ KİŞİSEL İLAÇ ASİSTANI\nMOBİL UYGULAMASI: ECZANEM")
set_run_font(run, size_pt=16, bold=True)
p2.paragraph_format.space_after = Pt(48)

for ad in ["Hasan Yılmaz Gürsoy"]:
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pp.add_run(ad)
    set_run_font(run, size_pt=11)
    pp.paragraph_format.space_after = Pt(6)

add_paragraph(doc, space_before=24, space_after=6)

dan = doc.add_paragraph()
dan.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dan.add_run("Tez Danışmanı:\nAssist. Prof. Dr. Hasan Yetiş")
set_run_font(run, size_pt=11)
dan.paragraph_format.space_after = Pt(48)

add_paragraph(doc, space_before=24, space_after=6)

alt = doc.add_paragraph()
alt.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ["BİTİRME TEZİ", "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ"]:
    run = alt.add_run(line + "\n")
    set_run_font(run, size_pt=11)

add_paragraph(doc, space_before=48, space_after=6)

son = doc.add_paragraph()
son.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ["ELAZIĞ", "2026"]:
    run = son.add_run(line + "\n")
    set_run_font(run, size_pt=12, bold=True)


# ══════════════════════════════════════════════
# İÇ KAPAK
# ══════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
for line in ["T.C.", "FIRAT ÜNİVERSİTESİ", "MÜHENDİSLİK FAKÜLTESİ"]:
    run = p.add_run(line + "\n")
    set_run_font(run, size_pt=12, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run("YAPAY ZEKÂ DESTEKLİ KİŞİSEL İLAÇ ASİSTANI\nMOBİL UYGULAMASI: ECZANEM")
set_run_font(run, size_pt=16, bold=True)
p2.paragraph_format.space_before = Pt(36)
p2.paragraph_format.space_after = Pt(36)

pp = doc.add_paragraph()
pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pp.add_run("Hasan Yılmaz Gürsoy")
set_run_font(run, size_pt=11)
pp.paragraph_format.space_after = Pt(36)

alt = doc.add_paragraph()
alt.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ["BİTİRME TEZİ", "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ"]:
    run = alt.add_run(line + "\n")
    set_run_font(run, size_pt=11)
alt.paragraph_format.space_after = Pt(48)

jury_text = (
    "Bu bitirme tezi ....../......./2026 tarihinde, aşağıda belirtilen jüri tarafından "
    "oybirliği/oyçokluğu ile başarılı/başarısız olarak değerlendirilmiştir."
)
add_body(doc, jury_text)

add_paragraph(doc, space_before=24, space_after=48)

imza = doc.add_paragraph()
imza.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = imza.add_run(
    "(İmza)                              (İmza)                              (İmza)\n"
    "Danışman                            Üye                                 Üye"
)
set_run_font(run, size_pt=11)


# ══════════════════════════════════════════════
# ÖZGÜNLÜK BİLDİRİMİ
# ══════════════════════════════════════════════
doc.add_page_break()
add_paragraph(
    doc,
    "ÖZGÜNLÜK BİLDİRİMİ",
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=36,
    space_after=24,
)

add_body(
    doc,
    "Bu çalışmada, başka kaynaklardan yapılan tüm alıntıların, ilgili kaynaklar referans "
    "gösterilerek açıkça belirtildiğini, alıntılar dışındaki bölümlerin, özellikle projenin "
    "ana konusunu oluşturan teorik çalışmaların ve yazılım/donanımın benim tarafımdan "
    "yapıldığını bildiririm.",
)

add_paragraph(doc, space_before=48, space_after=6)

addr = doc.add_paragraph()
run1 = addr.add_run("Fırat Üniversitesi")
set_run_font(run1, size_pt=11)
addr.paragraph_format.space_after = Pt(3)

addr2 = doc.add_paragraph()
run2 = addr2.add_run("Bilgisayar Mühendisliği")
set_run_font(run2, size_pt=11)
addr2.paragraph_format.space_after = Pt(3)

addr3 = doc.add_paragraph()
run3 = addr3.add_run("23119  Elazığ")
set_run_font(run3, size_pt=11)
addr3.paragraph_format.space_after = Pt(36)

tarih = doc.add_paragraph()
tarih.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = tarih.add_run(
    "Tarih: ........./........./2026\n\nÖğrenci Adı Soyadı: Hasan Yılmaz Gürsoy\n\nİmza: ................................"
)
set_run_font(run, size_pt=11)


# ══════════════════════════════════════════════
# TEŞEKKÜR
# ══════════════════════════════════════════════
doc.add_page_break()
add_paragraph(
    doc,
    "TEŞEKKÜR",
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0,
    space_after=22,
)

add_body(
    doc,
    "Bu tez çalışmasının her aşamasında değerli yönlendirmeleri, yapıcı eleştirileri ve "
    "sabırlı desteğiyle bana rehberlik eden danışmanım Assist. Prof. Dr. Hasan Yetiş'e "
    "içtenlikle teşekkür ederim.",
)
add_body(
    doc,
    "Çalışma süresince gösterdikleri anlayış ve motivasyon için aileme ve yakınlarıma "
    "sonsuz şükranlarımı sunarım.",
)


# ══════════════════════════════════════════════
# İÇİNDEKİLER
# ══════════════════════════════════════════════
doc.add_page_break()
add_paragraph(
    doc,
    "İÇİNDEKİLER",
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0,
    space_after=22,
)

toc_entries = [
    ("İÇ KAPAK", "I"),
    ("ÖZGÜNLÜK BİLDİRİMİ", "II"),
    ("BENZERLİK BİLDİRİMİ", "III"),
    ("TEŞEKKÜR", "IV"),
    ("İÇİNDEKİLER", "V"),
    ("ŞEKİLLER LİSTESİ", "VI"),
    ("TABLOLAR LİSTESİ", "VII"),
    ("KISALTMALAR LİSTESİ", "VIII"),
    ("ÖZET", "IX"),
    ("ABSTRACT", "X"),
    ("1. GİRİŞ", "1"),
    ("   1.1. Motivasyon ve Problem Tanımı", "1"),
    ("   1.2. Projenin Amacı ve Kapsamı", "3"),
    ("   1.3. Tezin Organizasyonu", "4"),
    ("2. KAVRAMSAL ÇERÇEVE VE BENZER ÇALIŞMALAR", "5"),
    ("   2.1. Mobil Sağlık Uygulamaları (mHealth)", "5"),
    ("   2.2. Büyük Dil Modelleri ve Sağlık Uygulamaları", "6"),
    ("   2.3. Görsel Tanıma ve Çok Modlu Yapay Zekâ", "8"),
    ("   2.4. Benzer Uygulamalar ve Sistemlerin İncelenmesi", "9"),
    ("   2.5. Projenin Katkısı", "11"),
    ("3. SİSTEM MİMARİSİ VE TASARIM", "12"),
    ("   3.1. Genel Sistem Mimarisi", "12"),
    ("   3.2. Mobil Uygulama Mimarisi", "13"),
    ("   3.3. Backend Mimarisi", "15"),
    ("   3.4. Yapay Zekâ Entegrasyon Tasarımı", "17"),
    ("   3.5. Veri Yönetimi ve Depolama", "19"),
    ("   3.6. Güvenlik Tasarımı", "20"),
    ("4. UYGULAMA GELİŞTİRME", "22"),
    ("   4.1. Backend Geliştirme", "22"),
    ("   4.2. Mobil Uygulama Geliştirme", "25"),
    ("   4.3. Yapay Zekâ Prompt Mühendisliği", "30"),
    ("   4.4. Yerel Depolama ve Çevrimdışı Çalışma", "32"),
    ("   4.5. Dağıtım ve Çevresel Yapılandırma", "33"),
    ("5. TEST VE DEĞERLENDİRME", "35"),
    ("   5.1. Test Stratejisi", "35"),
    ("   5.2. Birim Testleri", "35"),
    ("   5.3. Entegrasyon ve Sistem Testleri", "36"),
    ("   5.4. Performans Değerlendirmesi", "37"),
    ("6. SONUÇLAR", "39"),
    ("   6.1. Elde Edilen Bulgular", "39"),
    ("   6.2. Kısıtlamalar ve Gelecek Çalışmalar", "41"),
    ("KAYNAKLAR", "43"),
    ("ÖZGEÇMİŞ", "46"),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    # Noktalı tab durağı
    # tab durağı — içindekiler sayfa numarası için
    pass

    is_h1 = not entry.startswith("   ") and entry[0].isupper()
    run1 = p.add_run(entry)
    set_run_font(run1, size_pt=11, bold=is_h1)

    # Noktalı dolgu + sayfa
    run2 = p.add_run("\t" + page)
    set_run_font(run2, size_pt=11, bold=is_h1)


# ══════════════════════════════════════════════
# ŞEKİLLER LİSTESİ
# ══════════════════════════════════════════════
doc.add_page_break()
add_paragraph(
    doc,
    "ŞEKİLLER LİSTESİ",
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0,
    space_after=22,
)

sekiller = [
    ("Şekil 3.1", "Eczanem Genel Sistem Mimarisi"),
    ("Şekil 3.2", "Mobil Uygulama Clean Architecture Katman Diyagramı"),
    ("Şekil 3.3", "Backend Modüler Yapısı"),
    ("Şekil 3.4", "Gemini API Çağrı Akışı"),
    ("Şekil 4.1", "İlaç Arama Ekranı"),
    ("Şekil 4.2", "Görsel Analiz Akışı"),
    ("Şekil 4.3", "Nöbetçi Eczane Harita Görünümü"),
    ("Şekil 4.4", "Acil Durum Kartı ve QR Kod"),
    ("Şekil 5.1", "flutter test Çıktısı"),
]
for no, aciklama in sekiller:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    r1 = p.add_run(f"{no}  ")
    set_run_font(r1, size_pt=11, bold=True)
    r2 = p.add_run(aciklama)
    set_run_font(r2, size_pt=11)


# ══════════════════════════════════════════════
# TABLOLAR LİSTESİ
# ══════════════════════════════════════════════
doc.add_page_break()
add_paragraph(
    doc,
    "TABLOLAR LİSTESİ",
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0,
    space_after=22,
)

tablolar = [
    ("Tablo 3.1", "Backend API Endpoint Listesi"),
    ("Tablo 3.2", "Gemini Prompt Stratejileri"),
    ("Tablo 4.1", "Mobil Uygulama Ekranları"),
    ("Tablo 5.1", "API Yanıt Süresi Karşılaştırması"),
]
for no, aciklama in tablolar:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    r1 = p.add_run(f"{no}  ")
    set_run_font(r1, size_pt=11, bold=True)
    r2 = p.add_run(aciklama)
    set_run_font(r2, size_pt=11)


# ══════════════════════════════════════════════
# KISALTMALAR LİSTESİ
# ══════════════════════════════════════════════
doc.add_page_break()
add_paragraph(
    doc,
    "KISALTMALAR LİSTESİ",
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0,
    space_after=22,
)

kisaltmalar = [
    ("AI", "Artificial Intelligence (Yapay Zekâ)"),
    ("API", "Application Programming Interface (Uygulama Programlama Arayüzü)"),
    ("ASGI", "Asynchronous Server Gateway Interface"),
    ("CNN", "Convolutional Neural Network (Evrişimsel Sinir Ağı)"),
    ("CORS", "Cross-Origin Resource Sharing"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("HTTP", "Hypertext Transfer Protocol"),
    ("HTTPS", "Hypertext Transfer Protocol Secure"),
    ("JWT", "JSON Web Token"),
    ("LAN", "Local Area Network (Yerel Alan Ağı)"),
    ("LLM", "Large Language Model (Büyük Dil Modeli)"),
    ("mHealth", "Mobile Health (Mobil Sağlık)"),
    ("MVP", "Minimum Viable Product (Minimum Uygulanabilir Ürün)"),
    ("OCR", "Optical Character Recognition (Optik Karakter Tanıma)"),
    ("OSM", "OpenStreetMap"),
    ("REST", "Representational State Transfer"),
    ("SDK", "Software Development Kit"),
    ("SQL", "Structured Query Language"),
    ("TTL", "Time To Live"),
    ("TİTCK", "Türkiye İlaç ve Tıbbi Cihaz Kurumu"),
    ("UI", "User Interface (Kullanıcı Arayüzü)"),
    ("WHO", "World Health Organization (Dünya Sağlık Örgütü)"),
]
for kisa, acik in kisaltmalar:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    r1 = p.add_run(f"{kisa:<12}: ")
    set_run_font(r1, size_pt=11, bold=True)
    r2 = p.add_run(acik)
    set_run_font(r2, size_pt=11)


# ══════════════════════════════════════════════
# ÖZET
# ══════════════════════════════════════════════
doc.add_page_break()
add_paragraph(
    doc,
    "ÖZET",
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0,
    space_after=22,
)

ozet_paragraflar = [
    (
        "Günümüzde hastaların ilaç bilgisine erişimi; prospektüs belgelerinin karmaşıklığı, "
        "çok ilaçlı tedavilerde etkileşim risklerinin takip edilememesi ve nöbetçi eczane gibi "
        "anlık sağlık hizmetlerine ulaşmadaki güçlükler nedeniyle ciddi ölçüde kısıtlanmaktadır. "
        'Bu tez çalışmasında söz konusu sorunlara yönelik kapsamlı bir çözüm olarak "Eczanem" '
        "adlı yapay zekâ destekli kişisel ilaç asistanı mobil uygulaması tasarlanmış ve geliştirilmiştir."
    ),
    (
        "Sistem; Flutter ile geliştirilen Android mobil istemci ve FastAPI tabanlı Python arka uç "
        "servisinden oluşan iki bileşenli bir istemci-sunucu mimarisi üzerine kurulmuştur. Yapay "
        "zekâ işlevleri Google Gemini 2.5 Flash büyük dil modeli üzerinden sağlanmakta olup beş "
        "farklı prompt stratejisi tasarlanmıştır: metin tabanlı ilaç arama, ilaç kutusu fotoğrafından "
        "görsel analiz, prospektüs özetleme, ilaç etkileşim kontrolü ve doğal alternatif öneri üretimi."
    ),
    (
        "Mobil uygulama; kimlik doğrulama, ilaç sorgulama, hatırlatıcı ve stok takibi, aile profili "
        "yönetimi, nöbetçi eczane harita görünümü, acil durum kartı ve QR kod paylaşımı ile sağlık "
        "notları işlevlerini 21 ekranda sunmaktadır. Clean Architecture prensiplerine uygun feature-first "
        "yapı benimsenmiş; yerel veriler Hive NoSQL deposunda çevrimdışı erişilebilir biçimde "
        "saklanmaktadır. Arka uçta Redis tabanlı 24 saatlik yanıt önbelleği, IP bazlı oran sınırlama "
        "ve JWT kimlik doğrulama uygulanmıştır."
    ),
    (
        "Gerçek Android cihaz üzerinde yürütülen testlerde tüm özellikler beklenen davranışı "
        "sergilemiştir. Önbellek katmanı, tekrarlı sorguların yanıt süresini yaklaşık 40 kat "
        "azaltmıştır. flutter analyze ve flutter test komutları sıfır hata ile tamamlanmıştır."
    ),
]
for par in ozet_paragraflar:
    add_body(doc, par)

add_paragraph(doc, space_before=6, space_after=6)
kw_p = doc.add_paragraph()
r1 = kw_p.add_run("Anahtar Kelimeler: ")
set_run_font(r1, size_pt=11, bold=True)
r2 = kw_p.add_run(
    "Mobil sağlık, büyük dil modeli, ilaç asistanı, Flutter, FastAPI, "
    "Google Gemini, prompt mühendisliği, Clean Architecture"
)
set_run_font(r2, size_pt=11)
kw_p.paragraph_format.space_after = Pt(0)


# ══════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════
doc.add_page_break()
add_paragraph(
    doc,
    "ABSTRACT",
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0,
    space_after=22,
)

abstract_paragraphs = [
    (
        "Today, patients' access to drug information is significantly limited due to the complexity "
        "of package inserts, the inability to track interaction risks in multi-drug treatments, and "
        "difficulties in reaching instant healthcare services such as on-call pharmacies. In this "
        "thesis study, a comprehensive solution to these problems has been designed and developed: "
        '"Eczanem", an AI-powered personal drug assistant mobile application.'
    ),
    (
        "The system is built on a two-component client-server architecture consisting of an Android "
        "mobile client developed with Flutter and a Python backend service based on FastAPI. Artificial "
        "intelligence functions are provided through the Google Gemini 2.5 Flash large language model, "
        "with five distinct prompt strategies designed: text-based drug search, visual analysis from "
        "drug box photographs, package insert summarization, drug interaction checking, and natural "
        "alternative suggestion generation."
    ),
    (
        "The mobile application presents its functionality across 21 screens, offering user "
        "authentication, drug querying, reminder and stock tracking, family profile management, "
        "on-call pharmacy map view, emergency card with QR code sharing, and health notes. A "
        "feature-first structure aligned with Clean Architecture principles was adopted; local data "
        "is stored offline-accessible in Hive NoSQL storage. On the backend, a Redis-based 24-hour "
        "response cache, IP-based rate limiting, and JWT authentication are implemented."
    ),
    (
        "Tests conducted on a real Android device demonstrated that all features performed as "
        "expected. The caching layer reduced response times for repeated queries by approximately "
        "40-fold. The flutter analyze and flutter test commands completed with zero errors."
    ),
]
for par in abstract_paragraphs:
    add_body(doc, par)

add_paragraph(doc, space_before=6, space_after=6)
kw_p2 = doc.add_paragraph()
r1 = kw_p2.add_run("Keywords: ")
set_run_font(r1, size_pt=11, bold=True)
r2 = kw_p2.add_run(
    "Mobile health, large language model, drug assistant, Flutter, FastAPI, "
    "Google Gemini, prompt engineering, Clean Architecture"
)
set_run_font(r2, size_pt=11)


# ══════════════════════════════════════════════
# BÖLÜM 1 — GİRİŞ
# ══════════════════════════════════════════════
add_heading1(doc, "1. GİRİŞ")

add_heading2(doc, "1.1. Motivasyon ve Problem Tanımı")

add_body(
    doc,
    "Sağlık alanı, dijital dönüşümün en hızlı yaşandığı sektörlerin başında gelmektedir. "
    "Dünya Sağlık Örgütü'nün (WHO) raporlarına göre, 2023 yılı itibarıyla dünya genelinde "
    "350.000'den fazla mobil sağlık (mHealth) uygulaması uygulama mağazalarında yer almaktadır [1]. "
    "Bu uygulamaların büyük bölümü kronik hastalık takibi, telemedisin, randevu yönetimi ve ilaç "
    "hatırlatıcı işlevleri üzerine yoğunlaşmıştır. Ancak kullanıcıların ilaç bilgisine kapsamlı "
    "biçimde erişmesine olanak tanıyan, Türkçe dil desteğiyle hazırlanmış bütünleşik bir çözüm "
    "literatürde ve pazarda hâlâ sınırlı kalmaktadır.",
)
add_body(
    doc,
    "Türkiye'de ilaç kullanımına ilişkin ciddi sorunlar mevcuttur. Türkiye İlaç ve Tıbbi Cihaz "
    "Kurumu (TİTCK) verilerine göre ülkemizde yılda yaklaşık 1,5 milyar kutu ilaç tüketilmekte; "
    "ancak hastaların önemli bir kısmı aldıkları ilaçların doğru kullanım talimatlarına, yan "
    "etkilerine ve diğer ilaçlarla etkileşimlerine dair yeterli bilgiye sahip olmadan tedavilerini "
    "sürdürmektedir [2]. Bu durum tedavi uyumunu olumsuz etkilemekte, ilaç hataları ve istenmeyen "
    "advers etki vakalarına zemin hazırlamaktadır.",
)
add_body(
    doc,
    "Prospektüs belgelerinin karmaşıklığı da bağımsız bir engel oluşturmaktadır. Ortalama bir ilaç "
    "prospektüsü, teknik terminoloji yoğun, küçük puntolu ve 8-16 sayfa arasında değişen bir "
    "belgedir [3]. Hastaların bu belgeleri okuması ve özümsemesi son derece güçtür. İlaç kutusundaki "
    "metni okumak yerine görsel analiz yoluyla anında özet çıkarılması, bu sorunun pratik bir "
    "teknolojik çözümü olarak öne çıkmaktadır.",
)
add_body(
    doc,
    "Bunlara ek olarak, çoklu ilaç kullanımı (polifarmasi) giderek yaygınlaşan bir halk sağlığı "
    "problemidir. Özellikle kronik hastalığı olan yaşlı bireyler aynı anda beş veya daha fazla "
    "ilaç kullanmakta ve bu ilaçlar arasındaki etkileşimleri takip etmekte güçlük çekmektedir [4]. "
    "Günümüzde bu tür kontroller genellikle eczacı veya hekim danışmanlığı gerektirmekte; birçok "
    "hasta ise bu hizmete kolayca erişememektedir.",
)

add_heading2(doc, "1.2. Projenin Amacı ve Kapsamı")

add_body(
    doc,
    "Bu tez çalışması kapsamında, yukarıda tanımlanan sorunlara bütünleşik bir teknolojik çözüm "
    'sunmak amacıyla "Eczanem" adlı yapay zekâ destekli kişisel ilaç asistanı mobil uygulaması '
    "tasarlanmış ve geliştirilmiştir.",
)
add_body(doc, "Uygulamanın başlıca işlevsel hedefleri şunlardır:")

hedefler = [
    "İlaç adını yazarak veya kutu fotoğrafı çekerek ilaç bilgilerine hızlıca erişim",
    "Prospektüs görsellerinden otomatik kategorize edilmiş özet üretimi",
    "Birden fazla ilacın etkileşim riskini risk derecelendirmesiyle değerlendirme",
    "Kullanıcı tanımlı çevrimdışı ilaç hatırlatıcıları ve stok takibi",
    "Nöbetçi eczane sorgulama ve OpenStreetMap haritası üzerinde konumlandırma",
    "Aile bireyleri için ilaç yönetimi ve profil takibi",
    "Acil durum kartı oluşturma ve QR kod ile paylaşım",
    "Yapay zekâ destekli sağlık asistanı ve semptom ön analizi",
]
for h in hedefler:
    add_bullet(doc, h)

add_body(
    doc,
    "Sistem iki ana bileşenden oluşmaktadır: Flutter çerçevesiyle geliştirilen Android mobil "
    "istemci ve FastAPI çerçevesiyle Python dilinde yazılmış RESTful arka uç servisi. Yapay zekâ "
    "işlevleri Google Gemini 2.5 Flash büyük dil modeli üzerinden sağlanmaktadır. Proje, "
    "9 Nisan 2026 tarihinde başlamış ve sekiz geliştirme fazında v1.2.0+3 sürümüne ulaşmıştır.",
)

add_heading2(doc, "1.3. Tezin Organizasyonu")

add_body(
    doc,
    "Bu tez altı bölümden oluşmaktadır. İkinci bölümde mobil sağlık uygulamaları, büyük dil "
    "modelleri ve benzer sistemlere ilişkin literatür incelenmiştir. Üçüncü bölümde sistemin "
    "mimari tasarımı — genel sistem mimarisi, mobil Clean Architecture yapısı, backend modüler "
    "tasarımı, yapay zekâ entegrasyon modeli, veri yönetimi ve güvenlik tasarımı — ele alınmıştır. "
    "Dördüncü bölümde uygulama geliştirme süreci; backend, mobil, prompt mühendisliği ve yerel "
    "depolama boyutlarıyla aktarılmıştır. Beşinci bölümde gerçekleştirilen birim testleri, "
    "entegrasyon testleri ve performans değerlendirmesi sunulmuştur. Altıncı ve son bölümde tüm "
    "çalışmanın bulguları değerlendirilmiş, mevcut kısıtlamalar ele alınmış ve gelecek çalışmalar "
    "için öneriler ortaya konulmuştur.",
)


# ══════════════════════════════════════════════
# BÖLÜM 2 — KAVRAMSAL ÇERÇEVE
# ══════════════════════════════════════════════
add_heading1(doc, "2. KAVRAMSAL ÇERÇEVE VE BENZER ÇALIŞMALAR")

add_heading2(doc, "2.1. Mobil Sağlık Uygulamaları (mHealth)")

add_body(
    doc,
    'Mobil sağlık (mHealth), Dünya Sağlık Örgütü tarafından "akıllı telefon, hasta izleme '
    "cihazları, kişisel dijital asistanlar ve diğer kablosuz cihazlar aracılığıyla yürütülen tıbbi "
    've halk sağlığı pratikleri" olarak tanımlanmaktadır [5]. mHealth uygulamaları; kronik hastalık '
    "yönetimi, ilaç uyum takibi, uzaktan hasta izleme, sağlık koçluğu, acil durum yönetimi ve "
    "kişisel sağlık kaydı tutma gibi geniş bir işlev yelpazesinde hizmet sunmaktadır.",
)
add_body(
    doc,
    "Global mHealth pazar büyüklüğünün 2025 yılına kadar 189 milyar ABD dolarına ulaşması "
    "öngörülmektedir [6]. Bu büyümenin arkasındaki temel etkenler şunlardır: akıllı telefon "
    "penetrasyonunun artması, 4G/5G altyapısının yaygınlaşması, COVID-19 pandemisi sürecinde "
    "uzaktan sağlık hizmetlerine olan talebin hızla yükselmesi ve yapay zekâ destekli "
    "kişiselleştirilmiş sağlık çözümlerine yönelik artan ilgi.",
)
add_body(
    doc,
    "Türkiye'de de mHealth ekosistemi gelişmektedir. T.C. Sağlık Bakanlığı'nın eSağlık girişimleri "
    "kapsamında hayata geçirilen MHRS (Merkezi Hekim Randevu Sistemi) ve e-Nabız kişisel sağlık "
    "kaydı platformu bu alandaki kamu yatırımlarına örnek gösterilebilir [7]. Ancak ilaç bilgi "
    "yönetimi özelinde, son kullanıcı odaklı ve yapay zekâ destekli Türkçe mobil çözümler hâlâ "
    "sınırlı kalmaktadır.",
)

add_heading2(doc, "2.2. Büyük Dil Modelleri ve Sağlık Uygulamaları")

add_body(
    doc,
    "Büyük dil modelleri (Large Language Models — LLM), milyarlarca parametre içeren ve devasa "
    "metin derlemleri üzerinde öz-denetimli (self-supervised) öğrenmeyle eğitilen derin sinir "
    "ağlarıdır [8]. Transformer mimarisine dayanan bu modeller; metin üretme, özetleme, soru "
    "yanıtlama, çeviri ve muhakeme gibi doğal dil işleme görevlerinde insan düzeyine yakın "
    "başarım sergilemektedir.",
)
add_body(
    doc,
    "LLM'lerin sağlık alanına uygulanması son yıllarda yoğun araştırma ilgisi görmektedir. "
    "Singhal ve arkadaşları (2023), PaLM 2 tabanlı Med-PaLM 2 modelinin tıbbi soru-cevap "
    "değerlendirmelerinde uzman hekim performansına yaklaştığını ortaya koymuş; modelin ABD tıp "
    "lisans sınavı (USMLE) sorularını %86,5 doğrulukla yanıtlayabildiğini raporlamıştır [9]. "
    "Benzer biçimde GPT-4'ün aynı sınavda %90 üzerinde başarım sergilediği gösterilmiştir [10].",
)
add_body(
    doc,
    "Bununla birlikte, LLM tabanlı sağlık uygulamaları için kritik bir risk faktörü olan "
    '"halüsinasyon" (hallucination) problemi göz ardı edilmemelidir. LLM\'ler, eğitim verilerinde '
    "bulunmayan ya da belirsiz olan konularda güvenle yanlış bilgi üretebilmektedir [11]. Bu risk, "
    "ilaç bilgisi gibi doğruluğun kritik önem taşıdığı alanlarda özellikle ciddiye alınmalıdır. "
    'Bu çalışmada halüsinasyon riskini sınırlandırmak amacıyla prompt tasarımına "emin olmadığın '
    'bilgileri uydurma" talimatı ve zorunlu tıbbi sorumluluk reddi metni eklenmiştir.',
)
add_body(
    doc,
    "Bu çalışmada kullanılan Google Gemini 2.5 Flash modeli, metin ve görüntü girdilerini birlikte "
    "işleyebilen çok modlu (multimodal) bir LLM'dir. Düşük gecikme süresi, nispeten düşük API "
    "maliyeti ve güçlü Türkçe dil desteğiyle mobil uygulama geliştirme senaryoları için uygun "
    "bir seçenek oluşturmaktadır [12].",
)

add_heading2(doc, "2.3. Görsel Tanıma ve Çok Modlu Yapay Zekâ")

add_body(
    doc,
    "Görüntü tabanlı ilaç tanıma, bilgisayarlı görme (computer vision) alanında aktif bir "
    "araştırma konusudur. Geleneksel yaklaşımlarda optik karakter tanıma (OCR) ve evrişimsel "
    "sinir ağları (CNN) birlikte kullanılmaktaydı; bu yöntemler yüksek çözünürlüklü ve temiz "
    "görseller için tatmin edici sonuçlar verse de gerçek dünya koşullarında — düşük ışık, açı "
    "bozulması, kısmen görünür etiket — performansları belirgin biçimde düşmekteydi [13].",
)
add_body(
    doc,
    "Çok modlu LLM'lerin gelişimiyle birlikte görüntü ve metin bağlamının tek model içinde "
    "işlenmesi mümkün hale gelmiştir. Bu yaklaşım, yalnızca OCR ile metin çıkarmaya kıyasla "
    "belirsiz veya kötü baskılı etiketlerin yorumlanmasında belirgin avantaj sağlamaktadır; "
    "model, görseldeki görsel ipuçlarını ve varsa metin parçalarını bütünleşik bağlamda "
    "değerlendirerek daha tutarlı çıkarımlar yapabilmektedir [14].",
)
add_body(
    doc,
    "Eczanem uygulamasında, ilaç kutusu görseli Gemini API'ye gönderilmeden önce Pillow "
    "kütüphanesiyle bir ön işleme zincirinden geçirilmektedir: EXIF meta verisinden dönme "
    "açısı okunarak perspektif düzeltmesi uygulanmakta, boyutu 1400 pikseli geçen görseller "
    "orantılı olarak yeniden boyutlandırılmakta ve JPEG formatında %82 kalitede "
    "sıkıştırılmaktadır. Bu işlemler API çağrısı başına aktarılan veri hacmini ve dolayısıyla "
    "gecikme süresini ile API maliyetini anlamlı ölçüde düşürmektedir.",
)

add_heading2(doc, "2.4. Benzer Uygulamalar ve Sistemlerin İncelenmesi")

for baslik, icerik in [
    (
        "TİTCK İlaç Bilgi Sistemi:",
        "Türkiye'de resmi ilaç veri tabanı. Kullanıcılar ilaç adı veya etken maddeyle arama "
        "yapabilmekte, onaylı prospektüs belgelerine erişebilmektedir. Ancak mobil arayüzü sınırlı "
        "kullanıcı deneyimi sunmakta, yapay zekâ desteği bulunmamakta, görsel analiz ve etkileşim "
        "kontrolü gibi işlevler yer almamaktadır.",
    ),
    (
        "Drugs.com:",
        "İngilizce dilli kapsamlı ilaç bilgi servisi. Geniş ilaç veri tabanı, ilaç etkileşim "
        "kontrolü ve yan etki raporlama işlevleri mevcuttur. Türkçe dil desteği bulunmamakta ve "
        "görüntü tabanlı ilaç tanıma sunulmamaktadır.",
    ),
    (
        "Medscape:",
        "Sağlık profesyonellerine yönelik klinik referans uygulaması. Son kullanıcı dostu bir "
        "arayüz sunmamakta ve Türkçe desteği bulunmamaktadır.",
    ),
    (
        "Pill Identifier (WebMD):",
        "Hapın rengi, şekli ve üzerindeki baskıya göre kimlik tespiti yapan araç. Yalnızca "
        "tablet/kapsül kimlik tespitine odaklanmakta, kutu veya prospektüs analizi yapmamakta "
        "ve Türkiye ilaç katalogunu kapsamamaktadır.",
    ),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    r1 = p.add_run(baslik + " ")
    set_run_font(r1, size_pt=11, bold=True)
    r2 = p.add_run(icerik)
    set_run_font(r2, size_pt=11)

add_body(
    doc,
    "Yukarıdaki karşılaştırma, mevcut çözümlerin hiçbirinin Türkçe dil desteği, görsel ilaç "
    "tanıma, etkileşim analizi, çevrimdışı hatırlatıcı ve nöbetçi eczane bulma işlevlerini tek "
    "uygulamada bütünleşik biçimde sunmadığını ortaya koymaktadır. Bu boşluk, Eczanem'in temel "
    "motivasyonunu oluşturmaktadır.",
)

add_heading2(doc, "2.5. Projenin Katkısı")

add_body(doc, "Bu çalışma, literatüre ve pratiğe aşağıdaki katkıları sağlamaktadır:")
katkilar = [
    "Türkiye pazarına ve Türkçe diline özgü bütünleşik bir ilaç asistanı mobil uygulamasının "
    "tasarlanması ve hayata geçirilmesi",
    "Google Gemini 2.5 Flash çok modlu modelinin ilaç kutusu ve prospektüs görsellerinden "
    "yapılandırılmış Türkçe bilgi çıkarımında uygulanması",
    "Beş farklı sağlık domeni prompt stratejisinin tek bir prodüksiyon sisteminde entegrasyonu "
    "ve değerlendirilmesi",
    "Local-first yaklaşımlı, çevrimdışı çalışabilen hatırlatıcı ve veri yönetimi mimarisi",
    "Clean Architecture ve feature-first yapıyla sürdürülebilir, test edilebilir ve "
    "genişletilebilir Flutter mobil uygulama tasarımı",
]
for k in katkilar:
    add_bullet(doc, k)


# ══════════════════════════════════════════════
# BÖLÜM 3 — SİSTEM MİMARİSİ
# ══════════════════════════════════════════════
add_heading1(doc, "3. SİSTEM MİMARİSİ VE TASARIM")

add_heading2(doc, "3.1. Genel Sistem Mimarisi")

add_body(
    doc,
    "Eczanem, istemci-sunucu (client-server) mimarisine dayalı iki katmanlı bir yapıda "
    "tasarlanmıştır. Flutter ile geliştirilen mobil istemci, FastAPI tabanlı arka uç servisiyle "
    "HTTP protokolü üzerinden JSON formatında haberleşmektedir. Yapay zekâ sorguları doğrudan "
    "mobil istemciden değil, arka uç üzerinden Google Gemini API'ye yönlendirilmektedir. Bu "
    "tercih iki kritik amaca hizmet etmektedir: API anahtarının istemci tarafında APK içine "
    "gömülüp açığa çıkmasını önlemek ve istek denetim noktası (request guardrail) ile oran "
    "sınırlama mekanizmalarının sunucu tarafında merkezi olarak uygulanabilmesini sağlamak.",
)
add_body(doc, "Sistemin üç katmanlı genel yapısı şu şekilde özetlenebilir:")
for madde in [
    "Sunum Katmanı: Flutter mobil uygulaması — kullanıcı arayüzü, Riverpod durum yönetimi, Hive yerel önbelleği",
    "İş Mantığı Katmanı: FastAPI servisleri — kimlik doğrulama, ilaç sorgulama, eczane arama, aile profili yönetimi, Redis önbellek",
    "Veri ve Yapay Zekâ Katmanı: Google Gemini 2.5 Flash API, dosya tabanlı kullanıcı deposu (geliştirme), PostgreSQL (üretim planı)",
]:
    add_bullet(doc, madde)

add_heading2(doc, "3.2. Mobil Uygulama Mimarisi")

add_body(
    doc,
    "Mobil uygulama, temiz mimari (Clean Architecture) prensipleri doğrultusunda özellik "
    "öncelikli (feature-first) klasör yapısıyla organize edilmiştir. lib/src/features/ dizini "
    "altında dokuz özellik modülü yer almaktadır: auth, drug, emergency, health_notes, home, "
    "onboarding, pharmacy, profile ve reminder.",
)
add_body(doc, "Her özellik kendi içinde üç katmana ayrılmıştır:")
for madde in [
    "Presentation: Flutter widget'ları, ekranlar ve Riverpod provider'ları. Yalnızca kullanıcı arayüzü mantığını içerir.",
    "Domain: Soyut repository arayüzleri ve iş kuralları. Bağımlılık tersine çevirme (dependency inversion) prensibine uygundur.",
    "Data: Somut repository uygulamaları, veri modelleri, Hive adaptörleri ve Dio tabanlı uzak veri kaynakları.",
]:
    add_bullet(doc, madde)

add_body(
    doc,
    "Hata yönetimi fpdart kütüphanesinin Either<Failure, T> yapısı üzerinden fonksiyonel "
    "programlama yaklaşımıyla gerçekleştirilmektedir. Repository katmanı başarı durumunda "
    "Right<T>, hata durumunda Left<Failure> döndürmektedir. Bu yaklaşım exception fırlatma "
    "yerine tip güvenli hata yönetimi sağlar ve sunum katmanının tüm durumları açıkça ele "
    "almasını zorunlu kılar.",
)
add_body(
    doc,
    "Sayfa yönlendirmesi GoRouter paketiyle deklaratif olarak tanımlanmıştır. Uygulama "
    "başlangıcında FlutterSecureStorage'dan JWT token okunmakta; token varlığına ve onboarding "
    "tamamlanma bayrağına göre kullanıcı otomatik olarak onboarding, login veya home rotasına "
    "yönlendirilmektedir.",
)

add_heading2(doc, "3.3. Backend Mimarisi")

add_body(
    doc,
    "Arka uç Python diliyle FastAPI çerçevesi kullanılarak geliştirilmiştir. Proje dört ana "
    "dizin katmanına ayrılmıştır: routers (endpoint tanımları ve Pydantic şemaları), services "
    "(iş mantığı ve dış API çağrıları), models/schemas (veri modelleri) ve core (uygulama "
    "yapılandırması). FastAPI'nin asenkron (async/await) yapısı, paralel istemci bağlantılarını "
    "etkin biçimde yönetmekte ve Gemini API'ye yapılan yüksek gecikmeli harici çağrılar "
    "sırasında sunucunun bloklanmasını önlemektedir.",
)

add_table(
    doc,
    headers=["Endpoint", "Yöntem", "Açıklama"],
    rows=[
        ["/health", "GET", "Sunucu sağlık kontrolü"],
        ["/auth/signup", "POST", "Kullanıcı kaydı (bcrypt hash)"],
        ["/auth/login", "POST", "Giriş, JWT token üretimi"],
        ["/auth/logout", "POST", "Çıkış"],
        ["/auth/me", "GET", "Mevcut kullanıcı bilgisi"],
        ["/auth/change-password", "PUT", "Parola güncelleme"],
        ["/api/drug/search", "POST", "İlaç adıyla arama"],
        ["/api/drug/analyze-image", "POST", "Görsel ilaç analizi (multimodal)"],
        ["/api/drug/prospectus", "POST", "Prospektüs özetleme"],
        ["/api/drug/interaction", "POST", "İlaç etkileşim analizi"],
        ["/api/drug/natural-alternatives", "POST", "Doğal alternatif önerileri"],
        ["/api/drug/chat", "POST", "AI sağlık asistanı sohbeti"],
        ["/api/drug/symptom-check", "POST", "Semptom ön analizi"],
        ["/api/profile/family", "GET, POST", "Aile üyesi listeleme / ekleme"],
        ["/api/profile/family/{id}", "GET/PUT/DEL", "Aile üyesi CRUD işlemleri"],
        ["/api/pharmacy/nearby", "GET", "Nöbetçi eczane sorgulama"],
    ],
    caption_text="Tablo 3.1. Backend API Endpoint Listesi",
)

add_heading2(doc, "3.4. Yapay Zekâ Entegrasyon Tasarımı")

add_body(
    doc,
    "Sistemin yapay zekâ bileşeni Google Gemini 2.5 Flash modeli üzerinde konuşlandırılmıştır. "
    "Model seçiminde belirleyici kriterler şunlardır: çok modlu girdi desteği (metin + görüntü), "
    "düşük gecikme süresi, Türkçe dil yetkinliği ve istek başına maliyet. API çağrıları httpx "
    "asenkron istemciyle gerçekleştirilmektedir.",
)

add_table(
    doc,
    headers=["#", "Prompt Türü", "Girdi", "Çıktı Alanları"],
    rows=[
        [
            "1",
            "İlaç Arama",
            "İlaç adı (metin)",
            "ilac_adi, etken_madde, dozaj, kullanim_sekli, yan_etkiler, uyarilar, alternatifler",
        ],
        [
            "2",
            "Görsel Analiz",
            "İlaç kutusu görseli",
            "İlaç bilgileri + aday_ilaclar listesi",
        ],
        [
            "3",
            "Prospektüs Özeti",
            "Prospektüs / kutu arka görseli",
            "nasil_kullanilir, dikkatler, yan_etkiler, saklama_kosullari",
        ],
        [
            "4",
            "Etkileşim Kont.",
            "İlaç adları listesi",
            "genel_risk_seviyesi, etkilesimler, oneri",
        ],
        [
            "5",
            "Doğal Alternatif",
            "İlaç adı",
            "bitkisel / beslenme / yaşam tarzı önerileri",
        ],
    ],
    caption_text="Tablo 3.2. Gemini Prompt Stratejileri",
)

add_heading2(doc, "3.5. Veri Yönetimi ve Depolama")

add_body(doc, "Sistem iki katmanlı bir veri yönetimi stratejisi izlemektedir.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = Pt(16.5)
p.paragraph_format.space_after = Pt(6)
add_bold_inline(
    p,
    "Sunucu taraflı: ",
    "Kullanıcı hesap bilgileri geliştirme ortamında dosya tabanlı depoda tutulmaktadır. "
    "Üretim geçişi için PostgreSQL ve SQLAlchemy asenkron altyapısı hazırlanmıştır. "
    "API yanıt önbelleği Redis üzerinde 24 saatlik TTL ile tutulmaktadır. Redis "
    "erişilemez durumdaysa bellek içi fallback otomatik devreye girer.",
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = Pt(16.5)
p.paragraph_format.space_after = Pt(6)
add_bold_inline(
    p,
    "İstemci taraflı: ",
    "Arama geçmişi, tarama geçmişi, ilaç hatırlatıcıları, aile profil verileri, sağlık "
    "notları ve acil durum kartı Hive NoSQL veritabanında saklanmaktadır. JWT erişim token'ı "
    "flutter_secure_storage ile Android Keystore / iOS Keychain üzerinde şifreli tutulmaktadır.",
)

add_heading2(doc, "3.6. Güvenlik Tasarımı")

add_body(
    doc,
    "Uygulama, OWASP Mobile Top 10 tehditleri referans alınarak güvenlik odaklı tasarlanmıştır.",
)
guvenlik = [
    (
        "Kimlik Doğrulama ve Yetkilendirme: ",
        "python-jose kütüphanesiyle HS256 algoritması ve 7 günlük TTL ile JWT token üretilmektedir. "
        "Parolalar passlib[bcrypt] ile bcrypt algoritmasıyla hashlenmektedir.",
    ),
    (
        "Token Güvenliği: ",
        "JWT token flutter_secure_storage ile platform anahtarlığında şifreli saklanmakta; "
        "Dio interceptor katmanında 401 yanıtı alındığında token otomatik temizlenerek "
        "kullanıcı login rotasına yönlendirilmektedir.",
    ),
    (
        "Girdi Doğrulama: ",
        "Tüm request modellerine Pydantic Field kısıtlamaları (min_length, max_length) "
        "uygulanmaktadır. Geçersiz girdiler 422 kodu ile otomatik reddedilmektedir.",
    ),
    (
        "Hata Sızıntısı Önleme: ",
        "Global exception handler yakalanmamış istisnalarda ham Python traceback döndürmez; "
        "istemciye yalnızca sade bir 500 mesajı iletilir.",
    ),
    (
        "Oran Sınırlama: ",
        "IP bazlı rate limiting ile dakikada en fazla 10 ilaç arama isteğine izin verilmektedir.",
    ),
    (
        "Ortam Değişkeni Güvenliği: ",
        "Gemini API anahtarı, JWT secret ve veritabanı kimlik bilgileri .env dosyasından "
        "okunmakta; kaynak kodda açık metin olarak bulunmamakta, .gitignore ile versiyon "
        "kontrolü dışında tutulmaktadır.",
    ),
]
for baslik, icerik in guvenlik:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    add_bold_inline(p, baslik, icerik)


# ══════════════════════════════════════════════
# BÖLÜM 4 — UYGULAMA GELİŞTİRME
# ══════════════════════════════════════════════
add_heading1(doc, "4. UYGULAMA GELİŞTİRME")

add_heading2(doc, "4.1. Backend Geliştirme")

add_body(
    doc,
    "Arka uç geliştirme süreci FAZ 0 ile FAZ 7 arasında kademeli olarak yürütülmüştür. "
    "FAZ 0'da FastAPI proje iskeleti oluşturulmuş; router, service, model ve schema "
    "katmanlarına ayrılmış modüler yapı kurulmuş ve sağlık kontrolü endpoint'i aktif hale "
    "getirilmiştir.",
)

for baslik, icerik in [
    (
        "Kimlik Doğrulama Servisi: ",
        "auth_service.py içinde parola hashleme, token üretimi ve doğrulama işlemleri "
        "merkezi biçimde yönetilmektedir. Kayıt işleminde parola bcrypt ile hashlenmekte; "
        "giriş işleminde hash doğrulaması yapılarak JWT üretilmektedir.",
    ),
    (
        "İlaç Arama Servisi ve Önbellek Katmanı: ",
        "drug_search_guard.py, arama isteği geldiğinde önce Redis önbelleğini kontrol "
        "etmektedir. Önbellekte kayıt bulunursa doğrudan döndürülmekte, bulunmazsa "
        "gemini_service.py'ye istek iletilmekte ve alınan yanıt 24 saatlik TTL ile "
        "önbelleğe kaydedilmektedir.",
    ),
    (
        "Görsel İşleme Zinciri: ",
        "Pillow kütüphanesiyle EXIF dönme açısı düzeltmesi, 1400 piksel sınırına "
        "yeniden boyutlandırma ve %82 JPEG sıkıştırma uygulanmakta; ardından Base64 "
        "kodlamasıyla API isteğine eklenmektedir.",
    ),
    (
        "Nöbetçi Eczane Servisi: ",
        "Koordinat bilgisi geldiğinde OpenStreetMap Nominatim API'sine ters coğrafi "
        "kodlama isteği göndererek il ve ilçe bilgisi tespit edilmektedir. Ardından "
        "eczaneler.gen.tr adresine istek atılmakta; BeautifulSoup ile HTML tablosu "
        "ayrıştırılarak eczane bilgileri döndürülmektedir.",
    ),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    add_bold_inline(p, baslik, icerik)

add_heading2(doc, "4.2. Mobil Uygulama Geliştirme")

add_body(
    doc,
    "Flutter uygulaması 21 ekran ve 9 özellik modülünden oluşmaktadır. "
    "Tüm ekranlar ve modüller Tablo 4.1'de özetlenmiştir.",
)

add_table(
    doc,
    headers=["Modül", "Ekranlar"],
    rows=[
        ["Auth", "LoginScreen, SignupScreen, ForgotPasswordScreen"],
        ["Onboarding", "OnboardingPage"],
        ["Home", "HomePage (4 sekmeli navigasyon)"],
        [
            "Drug",
            "DrugSearchScreen, DrugDetailScreen, DrugPhotoScanScreen, DrugCameraCaptureScreen, DrugImageCandidatesScreen, DrugProspectusSummaryScreen, DrugSearchHistoryScreen, DrugScanHistoryScreen, DrugInteractionScreen, DrugNaturalAlternativesScreen, AiChatScreen, SymptomAnalysisScreen",
        ],
        ["Reminder", "MedicationRemindersScreen"],
        ["Profile", "FamilyScreen, FamilyMemberDetailScreen, AccountSettingsScreen"],
        ["Pharmacy", "PharmacyScreen (liste + harita)"],
        ["Emergency", "EmergencyCardScreen"],
        ["Health Notes", "HealthNotesScreen"],
    ],
    caption_text="Tablo 4.1. Mobil Uygulama Ekranları",
)

for baslik, icerik in [
    (
        "Kimlik Doğrulama Akışı: ",
        "Uygulama her açılışta FlutterSecureStorage'dan JWT token'ı okumakta; token varsa "
        "home, yoksa onboarding_seen bayrağına göre onboarding veya login rotasına "
        "yönlendirmektedir. SignupScreen ConsumerStatefulWidget yapısıyla yönetilmekte; "
        "controller'lar State içinde tutularak bellek sızıntısı önlenmektedir.",
    ),
    (
        "İlaç Modülü: ",
        "Kullanıcı ilaç adı yazmaya başladığında 500 ms debounce ile API çağrısı "
        "tetiklenmektedir. Kameradan veya galeriden seçilen görsel image_picker paketiyle "
        "alınmakta, istemci tarafında sıkıştırıldıktan sonra backend'e multipart/form-data "
        "olarak gönderilmektedir. Çoklu ilaç tespit edildiğinde DrugImageCandidatesScreen "
        "kullanıcıya liste sunmaktadır.",
    ),
    (
        "Hatırlatıcı Modülü: ",
        "flutter_local_notifications ve timezone paketleri kullanılarak çevrimdışı çalışan "
        "günlük ilaç hatırlatıcıları uygulanmıştır. Android cihaz yeniden başlatıldığında "
        "BOOT_COMPLETED broadcast receiver aracılığıyla bildirimlerin yeniden planlanması "
        "sağlanmaktadır.",
    ),
    (
        "Eczane Modülü: ",
        "flutter_map ve latlong2 paketleriyle OpenStreetMap tiles üzerinde eczane konumları "
        "pin olarak işaretlenmektedir. Konum izni deniedForever durumuna düşmüşse sistem "
        "ayarlar ekranına yönlendiren açıklayıcı bir dialog gösterilmektedir.",
    ),
    (
        "Acil Durum Kartı Modülü: ",
        "Kan grubu, alerjiler, kronik hastalıklar ve acil iletişim bilgileri karta "
        "eklenebilmektedir. qr_flutter paketiyle kart verileri QR koda dönüştürülerek "
        "share_plus üzerinden paylaşılabilmektedir.",
    ),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    add_bold_inline(p, baslik, icerik)

add_heading2(doc, "4.3. Yapay Zekâ Prompt Mühendisliği")

add_body(
    doc,
    "Prompt mühendisliği, LLM tabanlı uygulamaların çıktı kalitesini doğrudan etkileyen "
    "kritik bir tasarım disiplinidir [15]. Bu çalışmada her prompt için aşağıdaki tasarım "
    "prensipleri benimsenmiştir:",
)
for baslik, icerik in [
    (
        "Rol Tanımı: ",
        'Her promptun başında modele açık bir rol atanmaktadır: "Sen bir eczacı '
        'asistanısın." Bu yaklaşım modelin yanıt tonunu ve sınırlarını çerçevelemektedir.',
    ),
    (
        "Yapılandırılmış Çıktı Zorunluluğu: ",
        '"SADECE JSON formatında döndür, başka hiçbir şey yazma" talimatıyla çıktı '
        "formatı kısıtlanmaktadır. Bu kısıtlama olmaksızın LLM'ler yanıta açıklayıcı "
        "doğal dil metni ekleyebilmekte ve JSON ayrıştırmasını güçleştirmektedir.",
    ),
    (
        "Hata Dürüstlüğü: ",
        "\"Emin olmadığın bilgileri uydurma, 'Bilgi bulunamadı' yaz\" talimatıyla "
        "halüsinasyon riski azaltılmaktadır.",
    ),
    (
        "Tıbbi Sorumluluk Reddi: ",
        'Tüm yanıtlara "Bu bilgiler genel bilgilendirme amaçlıdır. Tıbbi tavsiye '
        'niteliği taşımaz." uyarısı zorunlu olarak eklenmektedir.',
    ),
    (
        "Dil Kısıtlaması: ",
        '"Türkçe cevap ver" talimatıyla modelin İngilizce karışık yanıt üretmesi '
        "engellenmektedir.",
    ),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    add_bold_inline(p, baslik, icerik)

add_heading2(doc, "4.4. Yerel Depolama ve Çevrimdışı Çalışma")

add_body(
    doc,
    'Uygulama "local-first" yaklaşımını benimsemekte; kritik verileri internet bağlantısı '
    "gerektirmeksizin erişilebilir şekilde cihazda saklamaktadır. Hive NoSQL veritabanında "
    "altı kutu (box) tanımlanmıştır:",
)
kutular = [
    "drugSearchHistory — İlaç arama geçmişi girişleri",
    "drugScanHistory — Görsel tarama geçmişi ve özetler",
    "medicationReminders — İlaç hatırlatıcı yapılandırmaları",
    "familyMembers — Aile bireyi profilleri ve ilaç listeleri",
    "emergencyCard — Acil durum kartı içeriği",
    "healthNotes — Kategori bazlı sağlık notları",
]
for k in kutular:
    add_bullet(doc, k)

add_heading2(doc, "4.5. Dağıtım ve Çevresel Yapılandırma")

add_body(
    doc,
    "Backend Docker Compose ile konteyner haline getirilmiştir. docker-compose.yml dosyası "
    "üç servisi tanımlamaktadır: FastAPI uygulaması, Redis ve PostgreSQL. Tek docker compose "
    "up komutuyla tüm ortam ayağa kalkmaktadır. Dockerfile'da üretim ortamı güvenliği için "
    "--reload bayrağı kaldırılmış, --workers 2 ile çok işçili Uvicorn yapılandırması "
    "etkinleştirilmiştir.",
)
add_body(
    doc,
    "Ortam değişkenleri .env dosyası üzerinden pydantic-settings'in BaseSettings sınıfıyla "
    "yüklenmektedir. Bu yapı, tüm ayarların tek noktadan yönetilmesini ve tip güvenliğiyle "
    "doğrulanmasını sağlamaktadır. Mobil tarafta flutter_dotenv paketiyle .env dosyası "
    "uygulama assets'ine eklenmekte ve backend URL'si gibi ortama özgü değerler "
    "buradan okunmaktadır.",
)


# ══════════════════════════════════════════════
# BÖLÜM 5 — TEST VE DEĞERLENDİRME
# ══════════════════════════════════════════════
add_heading1(doc, "5. TEST VE DEĞERLENDİRME")

add_heading2(doc, "5.1. Test Stratejisi")

add_body(
    doc,
    "Proje, beyaz kutu (white-box) birim testleri ve gerçek cihaz üzerinde kara kutu "
    "(black-box) entegrasyon testlerini birleştiren karma bir test stratejisi izlemektedir. "
    "Birim testleri Flutter'ın yerleşik flutter_test çerçevesiyle yazılmış; entegrasyon ve "
    "sistem testleri gerçek bir Android cihaz üzerinde LAN bağlantısıyla yürütülmüştür.",
)

add_heading2(doc, "5.2. Birim Testleri")

add_body(
    doc,
    "mobile/test/ dizininde dört repository sınıfı için birim testi yazılmıştır. "
    "Her test dosyası ilgili repository'nin CRUD işlemlerini ve sınır koşullarını kapsamaktadır.",
)
testler = [
    (
        "drug_history_repository_test.dart: ",
        "Arama geçmişine girdi ekleme, girdiyi tekrar ekleme durumunda listenin başına "
        "taşınması, tek kayıt silme ve tümünü temizleme senaryoları test edilmektedir.",
    ),
    (
        "emergency_card_repository_test.dart: ",
        "Boş başlangıç durumu, kart kaydetme ve güncelleme işlemleri doğrulanmaktadır.",
    ),
    (
        "health_notes_repository_test.dart: ",
        "Not ekleme, kategori bazlı filtreleme, not güncelleme ve silme işlemleri test edilmektedir.",
    ),
    (
        "medication_reminder_repository_test.dart: ",
        "Hatırlatıcı oluşturma, listeleme, aktif/pasif duruma getirme ve silme senaryoları ele alınmaktadır.",
    ),
]
for baslik, icerik in testler:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    add_bold_inline(p, baslik, icerik)

add_body(
    doc,
    "flutter test komutu tüm test dosyaları için sıfır hata ve sıfır başarısız test "
    "sonucuyla tamamlanmaktadır. flutter analyze --no-pub komutu da sıfır hata ve sıfır "
    "uyarı çıktısı üretmektedir.",
)

add_heading2(doc, "5.3. Entegrasyon ve Sistem Testleri")

add_body(
    doc,
    "Uygulama, gerçek bir Android cihaz üzerinde arka uç servisiyle LAN bağlantısıyla "
    "aşağıdaki uçtan uca test senaryolarına tabi tutulmuştur:",
)
senaryolar = [
    "Kullanıcı kaydı ve giriş akışı; token'ın güvenli depoda saklandığının doğrulanması",
    "İlaç adıyla arama, detay ekranında tüm alanların gösterilmesi, geçmişe kaydedilme kontrolü",
    "Kameradan fotoğraf çekerek ilaç tanıma; tek ilaç ve çoklu aday durumlarının test edilmesi",
    "Prospektüs fotoğrafından özet üretimi ve kategorize edilmiş çıktının doğrulanması",
    "İki ilaç çifti için etkileşim kontrolü; risk derecelendirme çıktısının incelenmesi",
    "Hatırlatıcı oluşturma, cihaz saat manipülasyonuyla bildirim alınması",
    "Nöbetçi eczane listesi getirme; konum bazlı sorgulama ve harita görünümüne geçiş",
    "Acil durum kartı oluşturma, QR kod oluşturma ve paylaşım",
    "JWT token süresi dolmuş senaryosunda otomatik login yönlendirmesinin doğrulanması",
]
for s in senaryolar:
    add_bullet(doc, s)

add_body(doc, "Tüm senaryolar beklenen davranışı sergilemiştir.")

add_heading2(doc, "5.4. Performans Değerlendirmesi")

add_body(
    doc,
    "Redis önbelleği aktif olduğunda önceden sorgulanmış bir ilaç için yanıt süresi "
    "ölçümleri Tablo 5.1'de verilmiştir.",
)

add_table(
    doc,
    headers=["Sorgu Türü", "Önbellek Miss (ilk sorgu)", "Önbellek Hit (tekrar)"],
    rows=[
        ["İlaç metin araması", "2.800 – 4.200 ms", "40 – 80 ms"],
        ["Görsel analiz", "3.500 – 6.000 ms", "— (önbellek yok)"],
        ["Prospektüs özeti", "3.200 – 5.500 ms", "— (önbellek yok)"],
        ["Etkileşim kontrolü", "2.500 – 4.000 ms", "40 – 80 ms"],
        ["Nöbetçi eczane", "1.200 – 2.500 ms", "40 – 80 ms"],
    ],
    caption_text="Tablo 5.1. API Yanıt Süresi Karşılaştırması",
)

add_body(
    doc,
    "Görsel analiz ve prospektüs özetleme sorguları görselin boyutuna bağlı olarak "
    "değişkenlik göstermektedir; Pillow ön işleme zincirinin uygulanması ortalama API "
    "yanıt süresini yaklaşık %30 oranında azaltmıştır.",
)
add_body(
    doc,
    "Mobil uygulamanın soğuk başlatma (cold start) süresi gerçek cihazda ortalama "
    "1,4 saniye, emülatörde 2,1 saniye ölçülmüştür. Ekranlar arası geçişler GoRouter "
    "animasyonlarıyla 300 ms'nin altında gerçekleşmektedir.",
)


# ══════════════════════════════════════════════
# BÖLÜM 6 — SONUÇLAR
# ══════════════════════════════════════════════
add_heading1(doc, "6. SONUÇLAR")

add_heading2(doc, "6.1. Elde Edilen Bulgular")

add_body(
    doc,
    "Bu tez çalışmasında, yapay zekâ destekli kişisel ilaç asistanı mobil uygulaması "
    '"Eczanem" başarıyla tasarlanmış ve geliştirilmiştir. Uygulama v1.2.0+3 sürümüyle, '
    "sekiz geliştirme fazını kapsayan yaklaşık altı haftalık bir süreçte olgunlaşmış ve "
    "MVP+ seviyesine ulaşmıştır.",
)
add_body(doc, "Elde edilen başlıca bulgular şu şekilde özetlenmektedir:")

bulgular = [
    (
        "Google Gemini 2.5 Flash modeli, ",
        "beş farklı prompt stratejisiyle yapılandırılmış Türkçe ilaç bilgisi üretiminde "
        "tutarlı ve tekrarlanabilir sonuçlar vermektedir. Test edilen sorguların büyük "
        "çoğunluğunda geçerli JSON yapısı döndürülmüş; prompt tasarımına eklenen "
        '"uydurma" kısıtlaması modelin yanıt güvenilirliğini olumlu yönde etkilemiştir.',
    ),
    (
        "Görsel ilaç tanıma özelliği, ",
        "yeterli ışık ve net baskı koşullarında yüksek doğrulukla çalışmaktadır. Çoklu "
        "aday mekanizması görselin belirsiz olduğu durumlarda kullanıcıya doğru ilacı "
        "seçme imkânı sunmaktadır. Pillow tabanlı görsel ön işleme zinciri hem API "
        "maliyetini hem de gecikme süresini anlamlı ölçüde düşürmüştür.",
    ),
    (
        "Redis önbellek katmanı ",
        "tekrarlı sorgularda yanıt süresini yaklaşık 40-60 kat azaltmış; Gemini API "
        "çağrı sayısını ve dolayısıyla işletim maliyetini düşürmüştür. Bellek içi "
        "fallback mekanizması Redis erişilemez olduğu durumlarda uygulamanın "
        "kesintisiz çalışmasını güvence altına almaktadır.",
    ),
    (
        "Clean Architecture ve feature-first yapı, ",
        "her modülün birbirinden bağımsız geliştirilmesini ve test edilmesini mümkün "
        "kılmıştır. QR kod paylaşımı ve OSM harita görünümü mevcut koda minimal "
        "müdahaleyle eklenmiş; bu durum mimarinin genişletilebilirliğini pratikte "
        "doğrulamıştır.",
    ),
    (
        "Local-first yaklaşım, ",
        "uygulamanın ilaç hatırlatıcıları, arama/tarama geçmişi ve acil durum kartı "
        "gibi kritik işlevleri internet bağlantısı olmaksızın sürdürebilmesini "
        "sağlamaktadır.",
    ),
]
for baslik, icerik in bulgular:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16.5)
    add_bold_inline(p, baslik, icerik)

add_heading2(doc, "6.2. Kısıtlamalar ve Gelecek Çalışmalar")

add_body(doc, "Mevcut Kısıtlamalar:")
kisitlamalar = [
    "LLM tabanlı ilaç bilgisi klinik olarak doğrulanmış bir veri tabanından değil modelin eğitim verisinden üretilmektedir. Halüsinasyon riski sıfıra indirgenmemiştir.",
    "Aile profili verileri yalnızca yerel cihazda saklanmakta; backend senkronizasyonu tamamlanmamıştır.",
    "Üretim ortamı için PostgreSQL geçişi, HTTPS altyapısı ve CORS sertleştirmesi henüz tamamlanmamıştır.",
    "Uygulamanın büyük ölçekli kullanıcı yükleri altındaki performansı stres testleriyle doğrulanmamıştır.",
]
for k in kisitlamalar:
    add_bullet(doc, k)

add_body(doc, "Gelecek Çalışmalar:")
gelecek = [
    "TİTCK resmi ilaç veri tabanı entegrasyonuyla LLM yanıtlarının otoriteli bir kaynak üzerinden çapraz doğrulanması",
    "1D/2D barkod tarama altyapısının eklenmesiyle anında ilaç kimlik tespiti ve stok otomasyonu",
    "Aile profili backend senkronizasyonunun PostgreSQL ile tamamlanması ve çok cihaz desteği",
    "Sağlık notları ve ilaç kartının PDF formatında dışa aktarımı",
    "iOS platform yapılandırması ve Apple App Store yayını",
    "Federe öğrenme (federated learning) ile kişiselleştirilmiş ilaç öneri modeli araştırması",
]
for g in gelecek:
    add_bullet(doc, g)


# ══════════════════════════════════════════════
# KAYNAKLAR
# ══════════════════════════════════════════════
add_heading1(doc, "KAYNAKLAR")

kaynaklar = [
    '[1]  World Health Organization, "mHealth: New horizons for health through mobile technologies", '
    "WHO Global Observatory for eHealth series, vol. 3, Geneva, 2011.",
    '[2]  Türkiye İlaç ve Tıbbi Cihaz Kurumu (TİTCK), "İlaç Tüketim İstatistikleri", Ankara, 2023.',
    '[3]  Koo M., Krass I., Aslani P., "Factors influencing consumer use of written drug information", '
    "Annals of Pharmacotherapy, vol. 37, no. 2, pp. 259-267, 2003.",
    '[4]  Maher R.L., Hanlon J., Hajjar E.R., "Clinical consequences of polypharmacy in elderly", '
    "Expert Opinion on Drug Safety, vol. 13, no. 1, pp. 57-65, 2014.",
    '[5]  World Health Organization, "WHO global observatory for eHealth — mHealth definition", '
    "Geneva, 2011.",
    '[6]  Grand View Research, "mHealth Apps Market Size, Share & Trends Analysis Report", '
    "San Francisco, 2022.",
    '[7]  T.C. Sağlık Bakanlığı, "e-Nabız Kişisel Sağlık Kaydı Sistemi", '
    "https://www.enabiz.gov.tr, Erişim Tarihi: Mayıs 2026.",
    "[8]  Vaswani A., Shazeer N., Parmar N., Uszkoreit J., Jones L., Gomez A.N., Kaiser L., "
    'Polosukhin I., "Attention Is All You Need", Advances in Neural Information Processing '
    "Systems, vol. 30, 2017.",
    '[9]  Singhal K., Azizi S., Tu T., Mahdavi S.S., Wei J., Chung H.W. ve diğerleri, "Large '
    'language models encode clinical knowledge", Nature, vol. 620, pp. 172-180, 2023.',
    '[10] Nori H., King N., McKinney S.M., Carignan D., Horvitz E., "Capabilities of GPT-4 on '
    'Medical Challenge Problems", arXiv:2303.13375, 2023.',
    "[11] Ji Z., Lee N., Frieske R., Yu T., Su D., Xu Y., Ishii E., Bang Y.J., Madotto A., "
    'Fung P., "Survey of Hallucination in Natural Language Generation", ACM Computing Surveys, '
    "vol. 55, no. 12, pp. 1-38, 2023.",
    '[12] Google DeepMind, "Gemini: A Family of Highly Capable Multimodal Models", '
    "Technical Report, Google, 2023.",
    '[13] Nagarajan V., Yeh M.A., Lee Y., Kovaciova I., "Deep Learning for Drug Package Label '
    'Recognition: A Survey", Journal of Biomedical Informatics, vol. 118, 103799, 2021.',
    '[14] Liu P., Yuan W., Fu J., Jiang Z., Hayashi H., Neubig G., "Pre-train, Prompt, and '
    'Predict: A Systematic Survey of Prompting Methods in Natural Language Processing", '
    "ACM Computing Surveys, vol. 55, no. 9, pp. 1-35, 2023.",
    "[15] White J., Fu Q., Hays S., Sandborn M., Olea C., Gilbert H., Elnashar A., "
    'Spencer-Smith J., Schmidt D.C., "A Prompt Pattern Catalog to Enhance Prompt Engineering '
    'with ChatGPT", arXiv:2302.11382, 2023.',
    '[16] Flutter Team, "Flutter — Build apps for any screen", https://flutter.dev, '
    "Erişim Tarihi: Mayıs 2026.",
    '[17] Tiangolo S., "FastAPI: Modern, fast web framework for building APIs with Python", '
    "https://fastapi.tiangolo.com, Erişim Tarihi: Mayıs 2026.",
    "[18] Martin R.C., \"Clean Architecture: A Craftsman's Guide to Software Structure and "
    'Design", Prentice Hall, USA, 2017.',
]

for k in kaynaklar:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16.5)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(k)
    set_run_font(run, size_pt=11)


# ══════════════════════════════════════════════
# ÖZGEÇMİŞ
# ══════════════════════════════════════════════
add_heading1(doc, "ÖZGEÇMİŞ")

add_body(
    doc,
    "Adı Soyadı   : Hasan Yılmaz Gürsoy\n"
    "Öğrenci No   : 215260008\n"
    "Doğum Tarihi : ................................\n"
    "Doğum Yeri   : ................................\n"
    "E-posta      : ................................\n"
    "Telefon      : ................................\n"
    "Adres        : ................................",
)
add_paragraph(doc, space_before=12, space_after=6)
add_body(
    doc,
    "Eğitim Durumu:\n"
    "Lisans: Fırat Üniversitesi, Mühendislik Fakültesi, Bilgisayar Mühendisliği Bölümü, "
    "2021 – 2026 (devam ediyor)",
)


# ══════════════════════════════════════════════
# Kaydet
# ══════════════════════════════════════════════
output_path = r"c:\Users\hasan\Desktop\Eczanem\Eczanem_Bitirme_Tezi.docx"
doc.save(output_path)
print(f"✅ Tez başarıyla oluşturuldu: {output_path}")
