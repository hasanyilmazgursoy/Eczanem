"""Vize ara raporunu Word (.docx) formatında üretir."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Sayfa kenar boşlukları
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Varsayılan stil
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

HEADING_COLOR = RGBColor(0x1A, 0x47, 0x6F)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADING_COLOR
    return h


def add_para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for par in row.cells[i].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def add_bold_bullet(label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(label).bold = True
    p.add_run(text)


def add_phase(title, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run(title).bold = True
    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.space_after = Pt(8)


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_section_label(text):
    """Kalın etiket satırı (örn. 'Bireysel Etki:')"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.add_run(text).bold = True


# ═══════════════════════════════════════════
# KAPAK SAYFASI
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("Eczanem — Bitirme Projesi\nVize Ara Raporu", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = HEADING_COLOR

doc.add_paragraph()
for label, value in [
    ("Öğrenci:", "Hasan"),
    ("Danışman:", "Hasan Yetiş"),
    ("Tarih:", "15 Nisan 2026"),
    ("Proje Başlangıç Tarihi:", "9 Nisan 2026"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = p.add_run(value)
    r2.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. GİRİŞ
# ═══════════════════════════════════════════
add_heading("1. Giriş ve Projenin Amacı")
add_para(
    "Eczanem; ilaç arama, fotoğraftan ilaç tanıma, prospektüs özetleme, ilaç etkileşim kontrolü, "
    "doğal alternatif önerisi ve ilaç hatırlatıcı gibi işlevleri tek bir mobil uygulamada birleştiren, "
    "yapay zekâ destekli bir kişisel ilaç asistanıdır. Flutter (Dart) ile geliştirilen mobil istemci "
    "ve FastAPI (Python) ile geliştirilen sunucu bileşenlerinden oluşmaktadır."
)
add_para(
    "Projenin temel motivasyonu, Türkiye'de hastaların ilaç bilgisine erişimde yaşadığı zorluklar, "
    "prospektüs metinlerinin karmaşıklığı ve kullanıcıların birden fazla ilacı bir arada kullanırken "
    "etkileşim risklerinden haberdar olamamasıdır. Eczanem, bu sorunlara teknolojik bir çözüm sunarak "
    "bireylerin sağlık okuryazarlığına katkı sağlamayı amaçlamaktadır."
)

# ═══════════════════════════════════════════
# 2. TEKNOLOJİLER VE YÖNTEMLER
# ═══════════════════════════════════════════
add_heading("2. Kullanılan Teknolojiler ve Yöntemler")

add_heading("2.1 Mimari Yaklaşım", level=2)
add_para("Proje, istemci-sunucu mimarisine dayalı, iki ana bileşenden oluşmaktadır:")
add_bold_bullet(
    "Mobil İstemci: ",
    "Flutter (Dart) ile geliştirilmiş, Clean Architecture prensiplerine uygun feature-first klasör "
    "yapısı. Her özellik kendi data, domain ve presentation katmanlarına sahiptir.",
)
add_bold_bullet(
    "Sunucu: ",
    "FastAPI (Python) ile geliştirilmiş, modüler yapıda (routers, services, models, schemas) RESTful API.",
)
add_bold_bullet(
    "Veri Akışı: ",
    "Mobil ↔ Backend arasında HTTPS üzerinden JSON tabanlı iletişim; yapay zekâ sorguları backend "
    "üzerinden Gemini API'ye yönlendirilmektedir.",
)
add_code_block(
    "Flutter Mobil Uygulama\n"
    " ├─ Auth (Kayıt / Giriş / Oturum)\n"
    " ├─ Home / 4 Sekmeli Navigasyon\n"
    " ├─ Drug Search (Metin + Görsel)\n"
    " ├─ Photo Scan + Prospectus Summary\n"
    " ├─ Search / Scan History\n"
    " ├─ Drug Interaction Check\n"
    " ├─ Natural Alternatives\n"
    " └─ Medication Reminder + Stock Tracking\n\n"
    "FastAPI Backend\n"
    " ├─ /health\n"
    " ├─ /api/auth/* (signup, login, me, logout)\n"
    " ├─ /api/drug/search\n"
    " ├─ /api/drug/analyze-image\n"
    " ├─ /api/drug/prospectus\n"
    " ├─ /api/drug/interaction\n"
    " └─ /api/drug/natural-alternatives"
)

add_heading("2.2 Mobil Teknoloji Yığını", level=2)
add_table(
    ["Teknoloji", "Kullanım Amacı"],
    [
        ["Flutter (Dart)", "Çapraz platform mobil uygulama geliştirme"],
        ["Riverpod", "Durum yönetimi (state management)"],
        ["Dio", "HTTP istemci katmanı"],
        ["GoRouter", "Deklaratif sayfa yönlendirme"],
        ["Hive", "Yerel NoSQL veri depolama"],
        ["flutter_secure_storage", "JWT token'ların şifrelenmiş saklanması"],
        ["camera / image_picker", "Kamera çekimi ve galeriden görsel seçme"],
        ["flutter_local_notifications", "Offline ilaç hatırlatıcı bildirimleri"],
        ["easy_localization", "Çoklu dil desteği (Türkçe / İngilizce)"],
        ["fpdart", "Fonksiyonel programlama (Either/Task pattern)"],
        ["Skeletonizer", "Yükleme durumu animasyonları"],
    ],
)

add_heading("2.3 Backend Teknoloji Yığını", level=2)
add_table(
    ["Teknoloji", "Kullanım Amacı"],
    [
        ["FastAPI (Python)", "Asenkron RESTful API sunucusu"],
        ["Google Gemini 2.5 Flash", "Yapay zekâ destekli ilaç analizi (metin + görsel)"],
        ["JWT + bcrypt", "Kimlik doğrulama ve parola güvenliği"],
        ["Redis", "API yanıt cache'leme (24 saat TTL) ve rate limiting"],
        ["Pillow (PIL)", "Görsel ön işleme (yeniden boyutlandırma, EXIF düzeltme)"],
        ["HTTPX", "Asenkron HTTP istemci (Gemini API çağrıları)"],
        ["Docker Compose", "Geliştirme ortamının konteynerizasyonu"],
    ],
)

add_heading("2.4 Yapay Zekâ Entegrasyonu — Google Gemini API", level=2)
add_para(
    "Projenin yapay zekâ bileşeni Google Gemini 2.5 Flash modeli üzerinden çalışmaktadır. "
    "Beş farklı prompt stratejisi tasarlanmış ve entegre edilmiştir:"
)
add_table(
    ["#", "Prompt Türü", "Giriş", "Çıkış"],
    [
        ["1", "İlaç Arama", "İlaç adı (metin)", "Etken madde, dozaj, kullanım, yan etkiler, uyarılar (JSON)"],
        ["2", "Görsel Analiz", "İlaç kutusu fotoğrafı", "Tanımlanan ilaç bilgileri + çoklu aday listesi"],
        ["3", "Prospektüs Özeti", "Prospektüs görseli", "Kategorize edilmiş özet (kullanım, dikkat, saklama vb.)"],
        ["4", "Etkileşim Kontrolü", "İlaç listesi", "Risk seviyesi (güvenli/dikkatli/tehlikeli) + detaylar"],
        ["5", "Doğal Alternatifler", "İlaç adı", "Bitkisel, beslenme ve yaşam tarzı önerileri"],
    ],
)
add_para(
    "Her prompt Türkçe ve yapılandırılmış JSON çıktı üretecek biçimde tasarlanmıştır. "
    "Multimodal çağrılarda görsel optimizasyonu (1400px max boyut, %82 JPEG sıkıştırma, "
    "EXIF döndürme düzeltmesi) uygulanarak API maliyeti ve yanıt süresi optimize edilmiştir."
)

add_heading("2.5 Veri Kaynakları", level=2)
add_para(
    "Proje, geleneksel anlamda statik bir veri seti kullanmamaktadır. Bunun yerine Gemini büyük "
    "dil modeli kendi eğitim verisi üzerinden yanıt üretmektedir. Ek veri kaynakları:"
)
add_bold_bullet(
    "Kullanıcı verileri: ",
    "Dosya tabanlı JSON deposu (geliştirme aşaması); üretim ortamı için PostgreSQL + SQLAlchemy altyapısı hazırlanmıştır.",
)
add_bold_bullet(
    "Yerel veri: ",
    "Hive NoSQL veritabanı üzerinden arama geçmişi, tarama geçmişi ve hatırlatıcı verileri cihazda saklanmaktadır.",
)
add_bold_bullet("Cache katmanı: ", "Redis tabanlı 24 saatlik sorgu cache'i ve bellek içi fallback mekanizması.")

add_heading("2.6 Güvenlik Önlemleri", level=2)
add_bold_bullet("Kimlik doğrulama: ", "JWT (HS256, 7 gün süre) + bcrypt parola hash'leme.")
add_bold_bullet(
    "Rate limiting: ",
    "IP bazlı kayar pencere algoritması ile dakikada 10 istek sınırı; cache hit'ler limite sayılmaz.",
)
add_bold_bullet("Token güvenliği: ", "Mobilde flutter_secure_storage ile şifrelenmiş depolama.")
add_bold_bullet(
    "Veri güvenliği: ",
    "Cache'ten çekilen veriler deepcopy() ile klonlanarak yan etki riski ortadan kaldırılmıştır.",
)
add_bold_bullet("Thread safety: ", "Kullanıcı deposu mutex (asyncio.Lock) ile korunmaktadır.")

doc.add_page_break()

# ═══════════════════════════════════════════
# 3. YAPILAN ÇALIŞMALAR
# ═══════════════════════════════════════════
add_heading("3. Yapılan Çalışmalar ve Özet Bulgular")
add_para(
    "Proje, 9 Nisan 2026 tarihinde başlamış olup bir haftalık yoğun geliştirme sürecinde "
    "aşağıdaki fazlar tamamlanmıştır:"
)

add_heading("3.1 Tamamlanan Modüller", level=2)

add_phase(
    "FAZ 0 — Altyapı ve Kurulum:",
    "Backend ve mobil proje iskeletleri oluşturulmuş, ortam değişkenleri, CORS ayarları, sağlık kontrolü "
    "endpoint'i, Docker Compose yapılandırması ve temel paket entegrasyonları tamamlanmıştır.",
)
add_phase(
    "FAZ 1 — Temel İlaç Sorgulama (MVP):",
    "Kullanıcılar ilaç adı yazarak arama yapabilmekte ve Gemini API üzerinden etken madde, dozaj, kullanım "
    "şekli, yan etkiler ve uyarı bilgilerini detaylı olarak görüntüleyebilmektedir. Arama debounce (500ms), "
    "skeleton loading animasyonları, arama geçmişinin yerel depolanması (8 kayıt limiti) ve kapsamlı hata "
    "yönetimi (internet yok, API hatası, boş sonuç) tamamlanmıştır. Redis tabanlı 24 saatlik cache ve IP "
    "bazlı rate limiting sayesinde tekrarlayan sorguların maliyeti sıfırlanmıştır.",
)
add_phase(
    "FAZ 2 — Kamera ile İlaç Tanıma ve Prospektüs Tarama:",
    "Kamera veya galeriden seçilen ilaç kutusu fotoğrafı, Gemini multimodal API ile analiz edilmektedir. "
    "Reçete veya blister fotoğrafında birden fazla ilaç tespit edildiğinde çoklu aday akışı sunulmaktadır. "
    "Prospektüs fotoğrafından kullanım, dikkat, saklama koşulları ve yan etkiler gibi başlıklara ayrılmış "
    "kategorize özet üretilmektedir. Görsel ön işleme (1400px max boyut, %82 JPEG sıkıştırma, EXIF "
    "döndürme) ile API maliyeti ve yanıt süresi optimize edilmiştir.",
)
add_phase(
    "Ara Faz — Geçmiş Merkezi ve Profil Kısayolları:",
    "Arama geçmişi ve tarama geçmişi ekranları (12 kayıt limiti, görsel/prospektüs modu ayrımı) "
    "oluşturularak profil sekmesindeki kısayollarla entegre edilmiştir. Geçmişten tekrar sorgulama, "
    "tekli silme ve tümünü temizleme işlevleri eklenmiştir.",
)
add_phase(
    "FAZ 4 — İlaç Hatırlatıcı ve Stok Takibi:",
    'Tamamen offline çalışan ilaç hatırlatıcı sistemi kurulmuştur. Kullanıcı belirli saatlerde hatırlatıcı '
    "kurabilmekte, günlük tekrar eden yerel bildirimler alabilmektedir. Stok takip dashboard'unda kalan "
    'tablet sayısı, ilerleme çubuğu ve otomatik bitiş süresi hesaplanmaktadır. "Aldım" butonu ile stok '
    "otomatik azalmakta, 3 gün önceden düşük stok uyarısı verilmektedir. Bildirimler cihaz yeniden "
    "başlatması sonrasında da otomatik olarak yeniden kurulmaktadır.",
)
add_phase(
    "FAZ 5 — İlaç Etkileşim Kontrolü ve Doğal Alternatifler:",
    "Kullanıcılar birden fazla ilacı girerek aralarındaki etkileşimi kontrol edebilmektedir. Sonuçlar renk "
    "kodlarıyla (Güvenli, Dikkatli, Tehlikeli) ve detaylı açıklamalarla sunulmaktadır. Her ilaç için bitkisel, "
    "beslenme ve yaşam tarzı bazında doğal alternatif önerileri de entegre edilmiştir.",
)

add_heading("3.2 Kısmen Tamamlanan Modüller", level=2)
add_phase(
    "FAZ 3 — Kullanıcı Sistemi ve Aile Profili:",
    "JWT tabanlı kayıt, giriş, oturum yönetimi ve bcrypt parola güvenliği tamamlanmıştır. "
    "Aile profili yönetimi (aile bireyi ekleme, birey bazlı ilaç listesi) henüz geliştirilmemiştir.",
)

add_heading("3.3 Özet Bulgular", level=2)
for f in [
    "Uygulama gerçek bir Android cihazda LAN bağlantısı üzerinden başarıyla test edilmiştir.",
    "Gemini API entegrasyonu beş farklı kullanım senaryosunda (metin arama, görsel analiz, prospektüs, etkileşim, doğal alternatif) kararlı çalışmaktadır.",
    "flutter analyze statik analiz kontrolü ve flutter test birim testleri sorunsuz geçmektedir.",
    "Redis cache katmanı sayesinde tekrarlayan ilaç sorgularında Gemini API maliyeti sıfıra düşmüştür.",
    "Rate limiting ile API'nin kötüye kullanımı engellenmiştir.",
    "Mobil uygulama 4 sekmeli ana sayfa, 10+ ekran ve offline çalışan hatırlatıcı desteğiyle MVP+ seviyesindedir.",
    "Toplam 6 backend endpoint, 10+ mobil ekran ve 2 test dosyası aktif olarak çalışmaktadır.",
]:
    add_bullet(f)

doc.add_page_break()

# ═══════════════════════════════════════════
# 4. FİNAL İÇİN PLAN
# ═══════════════════════════════════════════
add_heading("4. Final İçin Yapılabilecekler ve Beklenen Etki")

add_heading("4.1 Planlanan Geliştirmeler", level=2)
add_table(
    ["Faz", "Modül", "Açıklama", "Tahmini Süre"],
    [
        ["FAZ 3 (tamamlama)", "Aile Profili", "Aile bireyi ekleme, birey bazlı ilaç listesi, senkronizasyon", "1-2 hafta"],
        ["FAZ 6", "Nöbetçi Eczane + Sesli Sorgu", "Konum bazlı nöbetçi eczane, harita, Türkçe sesli arama", "1-2 hafta"],
        ["FAZ 7", "Acil Durum Kartı", "Kan grubu, alerji, kronik hastalık; PDF; sağlık günlüğü", "1 hafta"],
        ["FAZ 8", "Test ve Yayın", "UI polish, karanlık mod, onboarding, test, Play Store", "1-2 hafta"],
    ],
)

add_heading("4.2 Beklenen Etkiler", level=2)

add_section_label("Bireysel Etki:")
add_bullet("Kullanıcılar ilaç bilgilerine metin, fotoğraf veya ses yoluyla hızlı ve anlaşılır biçimde erişebilecek.")
add_bullet("İlaç etkileşim riskleri proaktif olarak tespit edilerek olası sağlık sorunlarının önüne geçilebilecek.")
add_bullet("Hatırlatıcı ve stok takip sistemiyle tedaviye uyum oranı artırılabilecek.")

add_section_label("Toplumsal Etki:")
add_bullet("Aile profili ile özellikle yaşlı bireyler ve çocuklar için merkezi ilaç takibi sağlanacak.")
add_bullet("Sesli sorgulama ile yaşlı veya görme engelli kullanıcılar da uygulamaya erişebilecek.")
add_bullet("Nöbetçi eczane entegrasyonu ile gece/tatil günü eczane arama sorunu çözülecek.")
add_bullet("Acil durum kartı kritik sağlık bilgilerini tek ekranda sunarak acil müdahaleyi hızlandıracak.")

add_section_label("Teknik ve Akademik Katkı:")
add_bullet("Büyük dil modellerinin (LLM) sağlık alanındaki pratik kullanımına somut bir örnek oluşturulacak.")
add_bullet("Multimodal yapay zekâ (metin + görsel) entegrasyonunun mobil sağlık uygulamalarındaki potansiyeli ortaya konacak.")
add_bullet("Flutter + FastAPI + Gemini API üçlüsünün Clean Architecture prensipleriyle bir arada kullanıldığı teknik bir referans proje sunulacak.")

# ═══════════════════════════════════════════
# 5. SONUÇ
# ═══════════════════════════════════════════
add_heading("5. Sonuç")
add_para(
    "Proje, 9 Nisan 2026'da başlamış olup bir haftalık yoğun geliştirme sürecinde sekiz fazlık yol "
    "haritasının beşi (FAZ 0, 1, 2, 4, 5) ve bir ara fazı başarıyla tamamlanmıştır. Uygulama; metin ve "
    "görsel tabanlı ilaç sorgulama, prospektüs özetleme, ilaç etkileşim kontrolü, doğal alternatif önerileri, "
    "offline ilaç hatırlatıcı ve stok takibi gibi çekirdek özelliklerle çalışan bir MVP+ seviyesindedir."
)
add_para(
    "Backend tarafında JWT kimlik doğrulama, Redis cache, rate limiting ve görsel optimizasyonu; "
    "mobil tarafta Clean Architecture, Riverpod durum yönetimi, Hive yerel depolama ve offline bildirim "
    "altyapısı kurulmuştur. Tüm modüller gerçek Android cihazda test edilmiştir."
)
add_para(
    "Final döneminde aile profili tamamlama, nöbetçi eczane entegrasyonu, sesli sorgulama, acil durum kartı "
    "ve Play Store yayını fazlarının gerçekleştirilmesi planlanmaktadır."
)

add_table(
    ["Gösterge", "Değer"],
    [
        ["Tamamlanan faz sayısı", "5 + 1 ara faz"],
        ["Kalan faz sayısı", "3 (+ 1 kısmi)"],
        ["Toplam backend endpoint", "6 aktif"],
        ["Toplam mobil ekran", "10+"],
        ["Yapay zekâ prompt türü", "5 farklı senaryo"],
        ["Test durumu", "flutter analyze ✓, flutter test ✓"],
        ["Gerçek cihaz doğrulaması", "Android (LAN üzerinden) ✓"],
    ],
)

# Sorumluluk reddi
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run(
    "Sorumluluk reddi: Bu uygulama genel bilgilendirme amaçlıdır ve tıbbi tavsiye niteliği "
    "taşımamaktadır. Kullanıcılar, ilaç kullanımıyla ilgili konularda mutlaka bir sağlık "
    "profesyoneline danışmalıdır."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── KAYDET ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Eczanem_Vize_Ara_Rapor.docx")
doc.save(output_path)
print(f"Rapor başarıyla oluşturuldu: {output_path}")
"""Vize ara raporunu Word (.docx) formatında üretir."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Sayfa kenar boşlukları
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Varsayılan stil
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    return h


def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Başlık satırı
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Veri satırları
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # Tablo sonrası boşluk
    return table


# ── KAPAK ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("Eczanem — Bitirme Projesi\nVize Ara Raporu", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)

doc.add_paragraph()
info_lines = [
    ("Öğrenci:", "Hasan"),
    ("Danışman:", "Hasan Yetiş"),
    ("Tarih:", "15 Nisan 2026"),
    ("Proje Başlangıç Tarihi:", "9 Nisan 2026"),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = p.add_run(f"{label} ")
    run_label.bold = True
    run_label.font.size = Pt(12)
    run_val = p.add_run(value)
    run_val.font.size = Pt(12)

doc.add_page_break()

# ── 1. PROJE TANIMI ──
add_heading("1. Proje Tanımı ve Amacı")
add_para(
    "Eczanem, kullanıcıların ilaç bilgilerini metin veya fotoğraf aracılığıyla sorgulayabildiği, "
    "ilaçlar arası etkileşim kontrolü yapabildiği, hatırlatıcı kurabildiği ve prospektüs özetleri "
    "oluşturabildiği kapsamlı bir mobil sağlık asistanı uygulamasıdır."
)
add_para(
    "Projenin temel amacı; bireylerin günlük ilaç kullanım süreçlerini kolaylaştırmak, ilaç bilgilerine "
    "hızlı ve anlaşılır biçimde erişim sağlamak ve olası ilaç etkileşimi risklerini kullanıcıya önceden "
    "bildirmektir. Uygulama, yapay zekâ destekli analiz yetenekleri sayesinde geleneksel ilaç rehberi "
    "uygulamalarının ötesine geçmeyi hedeflemektedir."
)

# ── 2. TEKNOLOJİLER ──
add_heading("2. Kullanılan Teknolojiler ve Yöntemler")

add_heading("2.1 Mobil Uygulama (İstemci)", level=2)
add_table(
    ["Teknoloji", "Kullanım Amacı"],
    [
        ["Flutter (Dart)", "Çapraz platform mobil uygulama geliştirme"],
        ["Riverpod", "State management (durum yönetimi)"],
        ["Dio", "HTTP istemci katmanı (backend iletişimi)"],
        ["GoRouter", "Deklaratif sayfa yönlendirme"],
        ["Hive", "Yerel veri depolama (geçmiş, hatırlatıcılar)"],
        ["flutter_secure_storage", "JWT token'ların güvenli saklanması"],
        ["camera / image_picker", "Kamera ile çekim ve galeriden görsel seçme"],
        [
            "flutter_local_notifications",
            "Offline çalışan ilaç hatırlatıcı bildirimleri",
        ],
        ["easy_localization", "Çoklu dil (Türkçe/İngilizce) desteği"],
    ],
)

add_heading("2.2 Backend (Sunucu)", level=2)
add_table(
    ["Teknoloji", "Kullanım Amacı"],
    [
        ["FastAPI (Python)", "RESTful API sunucusu"],
        ["Google Gemini API", "Yapay zekâ destekli ilaç sorgulama ve görsel analiz"],
        ["JWT", "Kullanıcı kimlik doğrulama ve oturum yönetimi"],
        ["Redis", "API yanıt cache'leme (24 saat) ve rate limiting"],
        ["bcrypt", "Parola hash'leme"],
        ["Docker Compose", "Geliştirme ortamı konteynerizasyonu"],
    ],
)

add_heading("2.3 Mimari Yaklaşım", level=2)
add_para(
    "Proje, Clean Architecture prensiplerine uygun olarak katmanlı bir yapıda tasarlanmıştır:"
)
bullets = [
    "Mobil taraf: Feature-first klasör yapısı (features/auth, features/drug, features/reminder); "
    "her özellik kendi data, domain ve presentation katmanlarına sahiptir.",
    "Backend taraf: Modüler FastAPI yapısı (routers, services, models, schemas) ile sorumluluklar ayrılmıştır.",
    "Veri akışı: Mobil ↔ Backend arasında HTTPS üzerinden JSON tabanlı RESTful iletişim; "
    "yapay zekâ sorguları backend üzerinden Gemini API'ye yönlendirilmektedir.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")
doc.add_paragraph()

add_heading("2.4 Yapay Zekâ Kullanımı", level=2)
add_para(
    "Projede Google Gemini büyük dil modeli (LLM), aşağıdaki görevler için kullanılmaktadır:"
)
ai_items = [
    "Metin tabanlı ilaç sorgulama: İlaç adı girilerek etken madde, dozaj, yan etkiler, uyarılar ve "
    "kullanım şekli bilgilerinin yapılandırılmış JSON formatında üretilmesi.",
    "Görsel ilaç tanıma: İlaç kutusu, blister veya etiket fotoğrafından ilacın tanımlanması (multimodal analiz).",
    "Prospektüs özetleme: Prospektüs veya kutu arkası görselinden okunabilir, kategorize edilmiş özet oluşturulması.",
    "İlaç etkileşim analizi: Birden fazla ilacın birlikte kullanım risklerinin değerlendirilmesi.",
    "Doğal alternatif önerileri: Belirli bir ilaç için bitkisel/doğal alternatif önerilerinin sunulması.",
]
for item in ai_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ── 3. YAPILANLAR ──
add_heading("3. Yapılan Çalışmalar ve Mevcut Durum")
add_para(
    "Proje, 9 Nisan 2026 tarihinde başlamış olup bugüne kadar aşağıdaki fazlar tamamlanmıştır:"
)

add_heading("3.1 Tamamlanan Fazlar", level=2)

phases = [
    (
        "FAZ 0 — Altyapı ve Kurulum",
        "Flutter proje iskeleti, backend API yapısı, ortam değişkenleri, CORS ayarları ve sağlık kontrolü "
        "endpoint'i hazırlanmıştır. Riverpod, Dio, GoRouter, Hive gibi temel paketler entegre edilmiştir.",
    ),
    (
        "FAZ 1 — Temel İlaç Sorgulama (MVP)",
        "Kullanıcı ilaç adını yazarak arama yapabilmekte, Gemini API üzerinden etken madde, dozaj, yan etkiler "
        "ve uyarı bilgilerini görüntüleyebilmektedir. Arama debounce (500ms), skeleton loading, arama geçmişi "
        "ve hata yönetimi tamamlanmıştır. Redis tabanlı 24 saat cache ve IP bazlı rate limiting aktiftir.",
    ),
    (
        "FAZ 2 — Kamera ile İlaç Tanıma ve Prospektüs Tarama",
        "Kamera veya galeriden seçilen ilaç kutusu fotoğrafı, Gemini multimodal API ile analiz edilmektedir. "
        "Çoklu ilaç adayı akışı (reçetede birden fazla ilaç tanıma) desteklenmektedir. Prospektüs fotoğrafından "
        "kategorize özet üretilmektedir. Görsel sıkıştırma ve yeniden boyutlandırma ile maliyet optimizasyonu yapılmıştır.",
    ),
    (
        "Ara Faz — Geçmiş Merkezi ve Profil Kısayolları",
        "Arama geçmişi ve tarama geçmişi ekranları oluşturulmuş, profil sekmesindeki kısayollarla entegre edilmiştir. "
        "Tekrar sorgulama, tekli silme ve tümünü temizleme işlevleri eklenmiştir.",
    ),
    (
        "FAZ 4 — İlaç Hatırlatıcı ve Stok Takibi",
        "Kullanıcı belirli saatlerde ilaç hatırlatıcısı kurabilmekte, günlük tekrar eden offline bildirimler "
        "alabilmektedir. Stok takip dashboard'u ile kalan tablet sayısına göre otomatik bitiş süresi hesaplanmakta, "
        "3 gün öncesinden düşük stok uyarısı verilmektedir. Bildirimler cihaz yeniden başlatması sonrası da korunmaktadır.",
    ),
    (
        "FAZ 5 — İlaç Etkileşim Kontrolü ve Doğal Alternatifler",
        "Kullanıcı birden fazla ilacı girerek aralarındaki etkileşimi kontrol edebilmektedir. Sonuçlar renk "
        "kodlarıyla (güvenli / dikkatli / tehlikeli) sunulmaktadır. İlaç detay ekranından doğal alternatif "
        "önerilerine erişim sağlanmaktadır.",
    ),
]
for phase_title, phase_desc in phases:
    add_heading(phase_title, level=3)
    add_para(phase_desc)

add_heading("3.2 Kısmen Tamamlanan Alanlar", level=2)
add_heading("FAZ 3 — Kullanıcı Sistemi ve Aile Profili", level=3)
add_para(
    "JWT tabanlı kayıt, giriş, oturum yönetimi ve şifre hash'leme altyapısı tamamlanmıştır. "
    "Aile profili yönetimi (aile bireyi ekleme, birey bazlı ilaç listesi) henüz geliştirilmemiştir."
)

add_heading("3.3 Çalışan Backend API Endpoint'leri", level=2)
add_table(
    ["Endpoint", "Yöntem", "İşlev"],
    [
        ["/health", "GET", "Sunucu sağlık kontrolü"],
        ["/api/auth/signup", "POST", "Kullanıcı kaydı"],
        ["/api/auth/login", "POST", "Giriş ve JWT token alma"],
        ["/api/auth/me", "GET", "Oturum bilgisi sorgulama"],
        ["/api/drug/search", "POST", "İlaç adıyla bilgi sorgulama"],
        ["/api/drug/analyze-image", "POST", "Görsel ilaç tanıma"],
        ["/api/drug/prospectus-summary", "POST", "Prospektüs özeti"],
        ["/api/drug/interactions", "POST", "İlaç etkileşim kontrolü"],
        ["/api/drug/natural-alternatives", "POST", "Doğal alternatif önerileri"],
    ],
)

add_heading("3.4 Doğrulama ve Test", level=2)
tests = [
    "Mobil uygulama flutter analyze ile statik analiz kontrolünden geçmektedir (sorun yok).",
    "flutter test ile hatırlatıcı repository ve ilaç geçmişi repository birim testleri başarıyla çalışmaktadır.",
    "Backend python -m compileall ile derleme doğrulamasından geçmektedir.",
    "Uygulama gerçek Android cihazda LAN bağlantısıyla test edilmiştir.",
]
for t in tests:
    doc.add_paragraph(t, style="List Bullet")

doc.add_page_break()

# ── 4. FİNAL ──
add_heading("4. Final İçin Yapılabilecekler ve Beklenen Etki")

add_heading("4.1 Planlanan Geliştirmeler", level=2)
add_table(
    ["Faz", "Modül", "Açıklama"],
    [
        [
            "FAZ 3",
            "Aile Profili",
            "Aile bireyi ekleme, birey bazlı ilaç listesi yönetimi",
        ],
        [
            "FAZ 6",
            "Nöbetçi Eczane + Sesli Sorgu",
            "Konum tabanlı nöbetçi eczane bulma, harita entegrasyonu, sesli ilaç arama",
        ],
        [
            "FAZ 7",
            "Acil Durum Kartı",
            "Kan grubu, alerji, kronik hastalık bilgileri; PDF dışa aktarım; sağlık günlüğü",
        ],
        [
            "FAZ 8",
            "Test ve Yayın",
            "UI polish, karanlık mod, onboarding, kapsamlı test yazımı, Play Store yayını",
        ],
    ],
)

add_heading("4.2 Beklenen Etkiler", level=2)
effects = [
    "Aile Profili ile uygulama bireysel kullanımın ötesine geçerek aile bazlı ilaç yönetimi sunacaktır. "
    "Özellikle yaşlı aile bireyleri veya çocuklar için ilaç takibini kolaylaştıracaktır.",
    "Nöbetçi Eczane entegrasyonu ile kullanıcılar en yakın açık eczaneyi harita üzerinde görebilecek, "
    "yol tarifi alabilecek ve doğrudan arayabilecektir.",
    "Sesli Sorgulama ile yaşlı veya görme engelli kullanıcılar da uygulamayı rahatça kullanabilecektir.",
    "Acil Durum Kartı kritik sağlık bilgilerini tek bir ekranda sunarak acil durumlarda hızlı bilgi erişimi sağlayacaktır.",
    "Play Store Yayını ile uygulama gerçek kullanıcılara ulaştırılabilecek, kullanılabilirlik geri bildirimleri toplanabilecektir.",
]
for e in effects:
    doc.add_paragraph(e, style="List Bullet")

add_heading("4.3 Projenin Genel Etkisi", level=2)
add_para(
    "Eczanem projesi, yapay zekâ tabanlı ilaç bilgi erişimi, görsel tanıma, etkileşim analizi ve hatırlatıcı "
    "sistemlerini tek bir mobil uygulamada birleştirmektedir. Proje tamamlandığında:"
)
impacts = [
    "Kullanıcılar ilaç bilgilerine metin, fotoğraf veya ses yoluyla hızlı erişim sağlayabilecek,",
    "İlaç etkileşim riskleri proaktif olarak tespit edilerek olası sağlık sorunlarının önüne geçilebilecek,",
    "Hatırlatıcı ve stok takip sistemiyle tedaviye uyum oranı artırılabilecek,",
    "Aile bazlı ilaç yönetimi ile birden fazla bireyin ilaç takibi merkezi bir noktadan yapılabilecektir.",
]
for imp in impacts:
    doc.add_paragraph(imp, style="List Bullet")

# ── 5. ÖZET ──
add_heading("5. Özet")
add_para(
    "Proje, 9 Nisan 2026'da başlamış olup bir haftalık süreçte planlanan 8 fazdan 5'i tamamlanmış, "
    "1'i kısmen hazırlanmıştır. Metin ve görsel tabanlı ilaç sorgulama, prospektüs özetleme, etkileşim "
    "kontrolü, doğal alternatif önerileri, ilaç hatırlatıcı ve stok takibi gibi çekirdek özellikler çalışır "
    "durumdadır. Backend ve mobil uygulama arasında güvenli iletişim, JWT kimlik doğrulama, Redis cache, "
    "rate limiting ve offline bildirim altyapısı kurulmuştur."
)
add_para(
    "Final dönemine kadar aile profili tamamlama, nöbetçi eczane entegrasyonu, sesli sorgulama, acil durum "
    "kartı ve uygulama yayınlama fazlarının gerçekleştirilmesi planlanmaktadır."
)

# Kaydet
output_path = os.path.join(os.path.dirname(__file__), "Eczanem_Vize_Ara_Rapor.docx")
doc.save(output_path)
print(f"Rapor oluşturuldu: {output_path}")
